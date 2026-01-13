import streamlit as st
import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime
import io
import re
import zipfile
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- 1. OCULTA A NAVEGAÇÃO PADRÃO ---
st.markdown(
    """
    <style>
        [data-testid="stSidebarNav"] {
            display: none;
        }
    </style>
    """,
    unsafe_allow_html=True
)

# --- 2. IMPORTAÇÃO E PROTEÇÃO DA PÁGINA ---
try:
    from Home import draw_sidebar, logout
except ImportError:
    def draw_sidebar():
        st.sidebar.error("Erro ao carregar a navegação. Execute a partir do Home.py.")
    def logout():
        st.sidebar.error("Erro ao carregar.")

st.session_state.current_page = 'Análise por Município'
draw_sidebar()

if not st.session_state.get('logged_in', False):
    st.error("Acesso negado. Por favor, faça o login na Página Principal.")
    st.page_link("Home.py", label="Ir para a página de Login", icon="🏠")
    st.stop()

# --- 3. INICIALIZAÇÃO DO SESSION STATE ---
if 'arquivos_gerados_municipio' not in st.session_state:
    st.session_state.arquivos_gerados_municipio = []

# --- CONFIGURAÇÕES GLOBAIS ---
requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)

MESES_MAPA = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
    "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}
LISTA_MESES = list(MESES_MAPA.keys())
meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# --- DICIONÁRIO DE MESORREGIÕES DE MG ---
MESORREGIOES_MG = {
    "Noroeste de Minas": ["Unaí", "Paracatu", "João Pinheiro", "Vazante", "Brasilândia de Minas", "Guarda-Mor", "Arinos", "Buritis"],
    "Norte de Minas": ["Montes Claros", "Janaúba", "Januária", "Pirapora", "Salinas", "Bocaiúva", "Porteirinha", "Jaíba", "Várzea da Palma", "Grão Mogol"],
    "Jequitinhonha": ["Diamantina", "Almenara", "Capelinha", "Araçuaí", "Minas Novas", "Itamarandiba", "Pedra Azul", "Jequitinhonha"],
    "Vale do Mucuri": ["Teófilo Otoni", "Nanuque", "Carlos Chagas", "Malacacheta", "Águas Formosas", "Itaipé"],
    "Triângulo Mineiro e Alto Paranaíba": ["Uberlândia", "Uberaba", "Araxá", "Patos de Minas", "Ituiutaba", "Frutal", "Patrocínio", "Araguari", "Conceição das Alagoas", "Sacramento", "Iturama"],
    "Central Mineira": ["Curvelo", "Três Marias", "Bom Despacho", "Felixlândia", "Corinto", "Pompeu", "Morada Nova de Minas"],
    "Metropolitana de Belo Horizonte": ["Belo Horizonte", "Betim", "Contagem", "Sete Lagoas", "Nova Lima", "Santa Luzia", "Ribeirão das Neves", "Ibirité", "Sabará", "Vespasiano", "Itabira", "Ouro Preto", "Mariana", "Congonhas", "Conselheiro Lafaiete"],
    "Vale do Rio Doce": ["Governador Valadares", "Ipatinga", "Coronel Fabriciano", "Timóteo", "Caratinga", "Aimorés", "Mantena", "Resplendor"],
    "Oeste de Minas": ["Divinópolis", "Formiga", "Itaúna", "Pará de Minas", "Nova Serrana", "Arcos", "Bambuí", "Piumhi", "Campo Belo", "Oliveira"],
    "Sul e Sudoeste de Minas": ["Poços de Caldas", "Pouso Alegre", "Varginha", "Passos", "Itajubá", "Alfenas", "Três Corações", "Lavras", "São Sebastião do Paraíso", "Guaxupé", "Extrema", "Varginha"],
    "Campos das Vertentes": ["Barbacena", "São João del Rei", "Lavras", "São Tiago", "Nazareno", "Barroso", "Resende Costa"],
    "Zona da Mata": ["Juiz de Fora", "Ubá", "Muriaé", "Manhuaçu", "Viçosa", "Cataguases", "Ponte Nova", "Leopoldina", "Santos Dumont", "Além Paraíba"]
}

# --- COLUNAS ---
MUN_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_MUN', 'CO_MUN', 'SH4', 'CO_SH4', 'CO_NCM']
MUN_DTYPES = {'CO_MUN': str, 'CO_SH4': str, 'SH4': str, 'CO_NCM': str, 'CO_PAIS': str}

# --- FUNÇÕES DE LÓGICA (Helpers) ---

def normalizar_codigo(codigo):
    """Remove .0 e espaços."""
    if pd.isna(codigo) or codigo == "": return None
    s = str(codigo).strip()
    if s.endswith('.0'): s = s[:-2]
    return s

def normalizar_pais(codigo):
    """Remove zeros à esquerda para padronizar países."""
    s = normalizar_codigo(codigo)
    if not s: return None
    if s.isdigit():
        return str(int(s))
    return s

def obter_lista_de_mesorregioes():
    return sorted(list(MESORREGIOES_MG.keys()))

def obter_municipios_da_meso(nome_meso):
    return MESORREGIOES_MG.get(nome_meso, [])

@st.cache_data(ttl=3600)
def ler_dados_csv_online(url, usecols=None, dtypes=None):
    retries = 3
    headers = {"User-Agent": "Mozilla/5.0"}
    for attempt in range(retries):
        try:
            resposta = requests.get(url, headers=headers, verify=False, timeout=(10, 1200)) 
            resposta.raise_for_status()
            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1', sep=';', dtype=dtypes)
            if 'CO_ANO' not in df.columns and '<!DOCTYPE' in str(df.columns): return None 
            if usecols:
                cols_existentes = [c for c in usecols if c in df.columns]
                df = df[cols_existentes]
            return df
        except (requests.exceptions.RequestException, ConnectionResetError) as e:
            if attempt < retries - 1:
                import time
                time.sleep(2)
                continue
            else: return None
        except Exception: return None
    return None

@st.cache_data(ttl=3600)
def carregar_dataframe(url, nome_arquivo, usecols=None, dtypes=None, mostrar_progresso=True):
    progress_bar = None
    if mostrar_progresso: progress_bar = st.progress(0, text=f"Carregando {nome_arquivo}...")
    df = ler_dados_csv_online(url, usecols=None, dtypes=dtypes)
    if mostrar_progresso and progress_bar: 
        if df is not None: progress_bar.progress(100, text=f"{nome_arquivo} carregado com sucesso.")
        else: progress_bar.empty()
    if df is not None and usecols:
         cols_to_keep = [c for c in usecols if c in df.columns]
         df = df[cols_to_keep]
    return df

@st.cache_data
def obter_dados_paises():
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv", usecols=['NO_PAIS', 'CO_PAIS'], dtypes={'CO_PAIS': str}, mostrar_progresso=False) 
    if df_pais is not None and not df_pais.empty:
        df_pais['CO_PAIS'] = df_pais['CO_PAIS'].apply(normalizar_pais)
        return pd.Series(df_pais.NO_PAIS.values, index=df_pais.CO_PAIS).to_dict()
    return {}

@st.cache_data
def obter_lista_de_municipios():
    url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
    df_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['SG_UF', 'NO_MUN', 'CO_MUN_GEO'], mostrar_progresso=False)
    if df_mun is not None:
        lista_mun = df_mun[df_mun['SG_UF'] == 'MG']['NO_MUN'].unique().tolist()
        lista_mun.sort()
        return lista_mun
    return ["Erro ao carregar"]

@st.cache_data
def obter_mapa_codigos_municipios():
    url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
    df_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['SG_UF', 'NO_MUN', 'CO_MUN_GEO'], dtypes={'CO_MUN_GEO': str}, mostrar_progresso=False)
    if df_mun is not None:
        df_mun_mg = df_mun[df_mun['SG_UF'] == 'MG'].copy()
        df_mun_mg['CO_MUN_GEO'] = df_mun_mg['CO_MUN_GEO'].apply(normalizar_codigo)
        return pd.Series(df_mun_mg.CO_MUN_GEO.values, index=df_mun_mg.NO_MUN).to_dict()
    return {}

@st.cache_data
def obter_dados_produtos_ncm():
    url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
    usecols_ncm = ['CO_SH2', 'NO_SH2_POR', 'CO_SH4', 'NO_SH4_POR']
    df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv", usecols=usecols_ncm, dtypes={'CO_SH4': str, 'CO_SH2': str}, mostrar_progresso=False)
    if df_ncm is not None:
        df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].apply(normalizar_codigo).str.zfill(4)
        df_ncm['CO_SH2_STR'] = df_ncm['CO_SH2'].apply(normalizar_codigo).str.zfill(2)
        mapa_sh4 = df_ncm.drop_duplicates('CO_SH4_STR').set_index('CO_SH4_STR')['NO_SH4_POR'].to_dict()
        mapa_sh2 = df_ncm.drop_duplicates('CO_SH2_STR').set_index('CO_SH2_STR')['NO_SH2_POR'].to_dict()
        return mapa_sh4, mapa_sh2
    return {}, {}

def get_sh4(codigo):
    s = normalizar_codigo(codigo)
    if not s: return None
    return s.zfill(4)[:4]

def get_sh2(sh4):
    s = normalizar_codigo(sh4)
    if s: return s[:2]
    return None

def normalizar_coluna_produto(df):
    if df is None: return None
    for col in ['CO_MUN', 'CO_PAIS', 'CO_SH4', 'CO_NCM']:
        if col in df.columns:
            df[col] = df[col].apply(normalizar_codigo)
    
    if 'CO_PAIS' in df.columns:
        df['CO_PAIS'] = df['CO_PAIS'].apply(normalizar_pais)

    if 'SH4' in df.columns: 
        df['SH4'] = df['SH4'].apply(get_sh4)
        return df
    if 'CO_SH4' in df.columns:
        df['SH4'] = df['CO_SH4'].apply(get_sh4)
    elif 'CO_NCM' in df.columns:
        df['SH4'] = df['CO_NCM'].apply(get_sh4)
    else:
        df['SH4'] = '0000'
    return df

def formatar_valor(valor):
    if pd.isna(valor): return "US$ 0,00"
    prefixo = ""
    if valor < 0: prefixo, valor = "-", abs(valor)
    if valor >= 1e9: return f"{prefixo}US$ {(valor/1e9):.2f} bilhões"
    if valor >= 1e6: return f"{prefixo}US$ {(valor/1e6):.2f} milhões"
    if valor >= 1e3: return f"{prefixo}US$ {(valor/1e3):.2f} mil"
    return f"{prefixo}US$ {valor:.2f}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def calc_var_display(row, col_atual, col_ant):
    v_atual = row[col_atual]
    v_ant = row[col_ant]
    if pd.isna(v_ant) or v_ant == 0:
        return "Novo Mercado" if v_atual > 0 else "-"
    var = ((v_atual - v_ant) / v_ant) * 100
    return f"{var:.2f}%"

def calcular_diferenca_percentual(valor_atual, valor_anterior):
    if pd.isna(valor_anterior) or valor_anterior == 0:
        return 100.0 if valor_atual > 0 else 0.0, "acréscimo" if valor_atual > 0 else "estabilidade"
    diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    tipo = "acréscimo" if diferenca > 0 else "redução" if diferenca < 0 else "estabilidade"
    return abs(diferenca), f"um {tipo}" if tipo != "estabilidade" else "uma estabilidade"

class DocumentoApp:
    def __init__(self, logo_path):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""
        self.logo_path = logo_path
        self.diretorio_base = "/tmp/" 
    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def adicionar_conteudo_formatado(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    def adicionar_titulo(self, texto):
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def nova_secao(self): pass
    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Cm(16.0))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(4.0)
        table.columns[1].width = Cm(12.0)
        cell_imagem = table.cell(0, 0)
        paragraph_imagem = cell_imagem.paragraphs[0]
        run_imagem = paragraph_imagem.add_run()
        if self.logo_path and os.path.exists(self.logo_path):
            try: run_imagem.add_picture(self.logo_path, width=Cm(3.5), height=Cm(3.42))
            except: pass
        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        for i, texto in enumerate(textos):
            p = cell_texto.paragraphs[0] if i == 0 else cell_texto.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = (i < 2)
    def finalizar_documento(self):
        try: os.makedirs(self.diretorio_base, exist_ok=True)
        except: pass
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue(), f"{sanitize_filename(self.titulo_doc)}.docx"

def clear_download_state_mun():
    st.session_state.arquivos_gerados_municipio = []

# Carregamento
lista_de_municipios = obter_lista_de_municipios()
mapa_codigos_municipios = obter_mapa_codigos_municipios()
mapa_nomes_paises = obter_dados_paises()
mapa_sh4_nomes, mapa_sh2_nomes = obter_dados_produtos_ncm()
lista_de_mesorregioes = obter_lista_de_mesorregioes()
ano_atual = datetime.now().year

# --- LAYOUT BALANCEADO (3 x 3) ---
col1, col2 = st.columns(2)
with col1:
    ano_principal = st.number_input("Ano de Referência:", min_value=1998, max_value=ano_atual, value=ano_atual, on_change=clear_download_state_mun)
    mesorregioes_selecionadas = st.multiselect("Filtrar por Mesorregião (opcional):", options=lista_de_mesorregioes, on_change=clear_download_state_mun)
    municipios_selecionados = st.multiselect("Selecione o(s) município(s):", options=lista_de_municipios, default=["BELO HORIZONTE"], on_change=clear_download_state_mun)

with col2:
    ano_comparacao = st.number_input("Ano de Comparação:", min_value=1998, max_value=ano_atual, value=ano_atual - 1, on_change=clear_download_state_mun)
    meses_selecionados = st.multiselect("Meses de Análise (opcional):", options=LISTA_MESES, on_change=clear_download_state_mun)
    top_n_itens = st.number_input("Nº de Itens nos Rankings:", min_value=1, max_value=100, value=10, on_change=clear_download_state_mun)

# --- Lógica de Agrupamento ---
if mesorregioes_selecionadas:
    municipios_da_meso = []
    for meso in mesorregioes_selecionadas:
        municipios_da_meso.extend(obter_municipios_da_meso(meso))
    todos_municipios = list(set(municipios_selecionados + municipios_da_meso))
else:
    todos_municipios = municipios_selecionados

agrupado = True
nome_agrupamento = None
nome_sugerido = ", ".join(mesorregioes_selecionadas) if (mesorregioes_selecionadas and not municipios_selecionados) else ""

if len(todos_municipios) > 1:
    st.header("2. Opções de Agrupamento")
    agrupamento_input = st.radio(f"Deseja que os dados dos {len(todos_municipios)} municípios sejam agrupados?", ("agrupados", "separados"), index=0, horizontal=True, on_change=clear_download_state_mun)
    agrupado = (agrupamento_input == "agrupados")
    if agrupado:
        quer_nome = st.checkbox("Dar nome ao agrupamento?", value=bool(nome_sugerido), on_change=clear_download_state_mun)
        if quer_nome:
            nome_agrupamento = st.text_input("Digite o nome:", value=nome_sugerido, on_change=clear_download_state_mun)
    st.header("3. Gerar Análise")
else:
    agrupado = False 
    st.header("2. Gerar Análise")

# --- EXECUÇÃO ---
if st.button("Iniciar Análise por Município"):
    st.session_state.arquivos_gerados_municipio = []
    logo_path_to_use = "LogoMinasGerais.png"
    
    with st.spinner(f"Processando {len(todos_municipios)} municípios..."):
        try:
            # Validação
            codigos_municipios_map = []
            municipios_validos = []
            for m in todos_municipios:
                cod = mapa_codigos_municipios.get(m) or mapa_codigos_municipios.get(m.upper())
                if cod:
                    codigos_municipios_map.append(normalizar_codigo(cod))
                    municipios_validos.append(m)
            
            if not codigos_municipios_map:
                st.error("Nenhum município válido.")
                st.stop()

            url_exp_mun_p = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_principal}_MUN.csv"
            url_exp_mun_c = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_comparacao}_MUN.csv"
            url_imp_mun_p = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_principal}_MUN.csv"
            url_imp_mun_c = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_comparacao}_MUN.csv"

            df_exp_mun_princ = carregar_dataframe(url_exp_mun_p, f"EXP_{ano_principal}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)
            df_exp_mun_comp = carregar_dataframe(url_exp_mun_c, f"EXP_{ano_comparacao}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)
            df_imp_mun_princ = carregar_dataframe(url_imp_mun_p, f"IMP_{ano_principal}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)
            df_imp_mun_comp = carregar_dataframe(url_imp_mun_c, f"IMP_{ano_comparacao}_MUN.csv", usecols=None, dtypes=MUN_DTYPES)

            if df_exp_mun_princ is None:
                st.error("Falha ao carregar dados.")
                st.stop()

            # Normalização
            df_exp_mun_princ = normalizar_coluna_produto(df_exp_mun_princ)
            df_exp_mun_comp = normalizar_coluna_produto(df_exp_mun_comp)
            df_imp_mun_princ = normalizar_coluna_produto(df_imp_mun_princ)
            df_imp_mun_comp = normalizar_coluna_produto(df_imp_mun_comp)

            for df in [df_exp_mun_princ, df_exp_mun_comp, df_imp_mun_princ, df_imp_mun_comp]:
                if df is not None:
                    df['SH2'] = df['SH4'].astype(str).str[:2]
                    if 'CO_MUN' in df.columns: df['CO_MUN'] = df['CO_MUN'].apply(normalizar_codigo)

            if meses_selecionados:
                meses_para_filtrar = [MESES_MAPA[m] for m in meses_selecionados]
                nome_periodo = f"o período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            else:
                meses_para_filtrar = list(range(1, df_exp_mun_princ['CO_MES'].max() + 1))
                nome_periodo = f"o ano de {ano_principal} (completo)"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"

            # CÁLCULO RANKING ESTADUAL
            def filtrar_mg_mes(df):
                return df[(df['SG_UF_MUN'] == 'MG') & (df['CO_MES'].isin(meses_para_filtrar))]

            df_exp_mg_total = filtrar_mg_mes(df_exp_mun_princ)
            df_imp_mg_total = filtrar_mg_mes(df_imp_mun_princ)
            
            total_exportacao_mg = df_exp_mg_total['VL_FOB'].sum()
            total_importacao_mg = df_imp_mg_total['VL_FOB'].sum()
            
            ranking_exp_mg = df_exp_mg_total.groupby('CO_MUN')['VL_FOB'].sum().sort_values(ascending=False)
            ranking_imp_mg = df_imp_mg_total.groupby('CO_MUN')['VL_FOB'].sum().sort_values(ascending=False)

            if not agrupado:
                municipios_para_processar = municipios_validos
            else:
                municipios_para_processar = [nome_agrupamento if (nome_agrupamento and nome_agrupamento.strip() != "") else ", ".join(municipios_validos)]

            for municipio_nome in municipios_para_processar:
                app = DocumentoApp(logo_path=logo_path_to_use)
                
                if agrupado:
                    st.subheader(f"Análise Agrupada: {municipio_nome}")
                    codigos_loop = codigos_municipios_map
                    nome_limpo = sanitize_filename(municipio_nome)
                    titulo_doc = f"Briefing - {nome_limpo} - {ano_principal}"
                    nome_doc = f"de {municipio_nome}"
                    posicao_exp_mg = "-"
                    part_exp_mg = 0
                    posicao_imp_mg = "-"
                    part_imp_mg = 0
                else:
                    st.subheader(f"Análise: {municipio_nome}")
                    c = mapa_codigos_municipios.get(municipio_nome) or mapa_codigos_municipios.get(municipio_nome.upper())
                    cod_mun = normalizar_codigo(c)
                    codigos_loop = [cod_mun]
                    nome_limpo = sanitize_filename(municipio_nome)
                    titulo_doc = f"Briefing - {nome_limpo} - {ano_principal}"
                    nome_doc = f"de {municipio_nome}"
                    
                    try: posicao_exp_mg = ranking_exp_mg.index.get_loc(cod_mun) + 1
                    except: posicao_exp_mg = "-"
                    try: posicao_imp_mg = ranking_imp_mg.index.get_loc(cod_mun) + 1
                    except: posicao_imp_mg = "-"

                app.set_titulo(titulo_doc)

                # FLUXO
                df_exp_princ_f = df_exp_mun_princ[(df_exp_mun_princ['CO_MUN'].isin(codigos_loop)) & (df_exp_mun_princ['CO_MES'].isin(meses_para_filtrar))]
                df_exp_comp_f = df_exp_mun_comp[(df_exp_mun_comp['CO_MUN'].isin(codigos_loop)) & (df_exp_mun_comp['CO_MES'].isin(meses_para_filtrar))]
                df_imp_princ_f = df_imp_mun_princ[(df_imp_mun_princ['CO_MUN'].isin(codigos_loop)) & (df_imp_mun_princ['CO_MES'].isin(meses_para_filtrar))]
                df_imp_comp_f = df_imp_mun_comp[(df_imp_mun_comp['CO_MUN'].isin(codigos_loop)) & (df_imp_mun_comp['CO_MES'].isin(meses_para_filtrar))]
                
                val_exp = df_exp_princ_f['VL_FOB'].sum()
                val_exp_ant = df_exp_comp_f['VL_FOB'].sum()
                val_imp = df_imp_princ_f['VL_FOB'].sum()
                val_imp_ant = df_imp_comp_f['VL_FOB'].sum()
                
                fluxo = val_exp + val_imp
                fluxo_ant = val_exp_ant + val_imp_ant
                saldo = val_exp - val_imp
                saldo_ant = val_exp_ant - val_imp_ant
                
                dif_fluxo, tipo_fluxo = calcular_diferenca_percentual(fluxo, fluxo_ant)
                
                app.nova_secao()
                app.adicionar_titulo("1. Fluxo Comercial")
                texto_fluxo = (f"Em {ano_principal}, {nome_doc} teve um fluxo comercial de {formatar_valor(fluxo)}, "
                               f"representando {tipo_fluxo} de {dif_fluxo:.1f}% em comparação a {ano_comparacao}. "
                               f"A balança comercial fechou em {formatar_valor(saldo)}.")
                app.adicionar_conteudo_formatado(texto_fluxo)

                # EXPORTAÇÃO
                dif_exp, tipo_exp = calcular_diferenca_percentual(val_exp, val_exp_ant)
                part_exp = (val_exp / total_exportacao_mg * 100) if total_exportacao_mg else 0
                
                app.nova_secao()
                app.adicionar_titulo("2. Exportações")
                texto_exp_1 = (f"As exportações {nome_doc} somaram {formatar_valor(val_exp)} em {ano_principal}, "
                               f"representando {tipo_exp} de {dif_exp:.1f}% em comparação a {ano_comparacao}.")
                app.adicionar_conteudo_formatado(texto_exp_1)
                
                if not agrupado:
                    texto_exp_2 = (f"{municipio_nome} foi o {posicao_exp_mg}º principal município exportador de Minas Gerais em {ano_principal}, "
                                   f"com uma participação de {part_exp:.2f}% nas vendas de Minas.")
                    app.adicionar_conteudo_formatado(texto_exp_2)
                
                exp_paises = df_exp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False).head(5)
                exp_prods = df_exp_princ_f.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(5)
                
                top_paises_txt = []
                for c, v in exp_paises.items():
                    nm = mapa_nomes_paises.get(normalizar_pais(c), f"Desconhecido ({c})")
                    pc = (v/val_exp)*100 if val_exp else 0
                    top_paises_txt.append(f"{nm} ({pc:.1f}%)")
                
                top_prods_txt = []
                for c, v in exp_prods.items():
                    nm = mapa_sh4_nomes.get(c, c)
                    pc = (v/val_exp)*100 if val_exp else 0
                    top_prods_txt.append(f"{nm} ({pc:.1f}%)")

                if top_paises_txt:
                    app.adicionar_conteudo_formatado(f"Principais destinos: {'; '.join(top_paises_txt)}.")
                if top_prods_txt:
                    app.adicionar_conteudo_formatado(f"Principais produtos: {'; '.join(top_prods_txt)}.")

                # VISUAL EXP
                st.header(f"Exportações")
                st.subheader("Principais Destinos")
                exp_p_p = df_exp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                exp_p_c = df_exp_comp_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                exp_p_p['País'] = exp_p_p['CO_PAIS'].apply(normalizar_pais).map(mapa_nomes_paises).fillna("Desconhecido")
                exp_p_c['País'] = exp_p_c['CO_PAIS'].apply(normalizar_pais).map(mapa_nomes_paises).fillna("Desconhecido")
                
                exp_final = pd.merge(exp_p_p, exp_p_c, on='País', how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                col_princ = f'VL_FOB {ano_principal}'
                col_comp = f'VL_FOB {ano_comparacao}'
                exp_final['Variação %'] = exp_final.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                exp_final = exp_final.sort_values(by=col_princ, ascending=False)
                exp_final = exp_final.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}'})
                
                df_show = exp_final.copy()
                df_show[f'Valor {ano_principal}'] = df_show[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show[f'Valor {ano_comparacao}'] = df_show[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show.head(top_n_itens), hide_index=True, use_container_width=True)

                st.subheader("Principais Produtos")
                exp_pr_p = df_exp_princ_f.groupby('SH4')['VL_FOB'].sum().reset_index()
                exp_pr_c = df_exp_comp_f.groupby('SH4')['VL_FOB'].sum().reset_index()
                exp_pr_p['Descrição'] = exp_pr_p['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                
                exp_final_pr = pd.merge(exp_pr_p, exp_pr_c, on='SH4', how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                exp_final_pr['Variação %'] = exp_final_pr.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                exp_final_pr = exp_final_pr.sort_values(by=col_princ, ascending=False)
                
                # Preenche descrição perdida no merge
                exp_final_pr['Descrição'] = exp_final_pr['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                
                exp_final_pr = exp_final_pr.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}', 'SH4': 'Código SH4'})
                
                df_show_pr = exp_final_pr.copy()
                df_show_pr[f'Valor {ano_principal}'] = df_show_pr[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_pr[f'Valor {ano_comparacao}'] = df_show_pr[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_pr[['Código SH4', 'Descrição', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_itens), hide_index=True, use_container_width=True)

                # IMPORTAÇÃO
                dif_imp, tipo_imp = calcular_diferenca_percentual(val_imp, val_imp_ant)
                part_imp = (val_imp / total_importacao_mg * 100) if total_importacao_mg else 0
                
                app.nova_secao()
                app.adicionar_titulo("3. Importações")
                texto_imp_1 = (f"As importações {nome_doc} somaram {formatar_valor(val_imp)} em {ano_principal}, "
                               f"representando {tipo_imp} de {dif_imp:.1f}% em comparação a {ano_comparacao}.")
                app.adicionar_conteudo_formatado(texto_imp_1)
                
                if not agrupado:
                    texto_imp_2 = (f"{municipio_nome} foi o {posicao_imp_mg}º principal município importador de Minas Gerais em {ano_principal}, "
                                   f"com uma participação de {part_imp:.2f}% nas compras de Minas.")
                    app.adicionar_conteudo_formatado(texto_imp_2)

                # Texto IMP
                imp_paises = df_imp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False).head(5)
                imp_prods = df_imp_princ_f.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(5)

                top_paises_imp_txt = []
                for c, v in imp_paises.items():
                    nm = mapa_nomes_paises.get(normalizar_pais(c), "Desconhecido")
                    pc = (v/val_imp)*100 if val_imp else 0
                    top_paises_imp_txt.append(f"{nm} ({pc:.1f}%)")
                
                top_prods_imp_txt = []
                for c, v in imp_prods.items():
                    nm = mapa_sh4_nomes.get(c, c)
                    pc = (v/val_imp)*100 if val_imp else 0
                    top_prods_imp_txt.append(f"{nm} ({pc:.1f}%)")

                if top_paises_imp_txt:
                    app.adicionar_conteudo_formatado(f"Principais origens: {'; '.join(top_paises_imp_txt)}.")
                if top_prods_imp_txt:
                    app.adicionar_conteudo_formatado(f"Principais produtos: {'; '.join(top_prods_imp_txt)}.")

                # Visual IMP
                st.header(f"Importações")
                st.subheader("Principais Origens")
                
                imp_p_p = df_imp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                imp_p_c = df_imp_comp_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                imp_p_p['País'] = imp_p_p['CO_PAIS'].apply(normalizar_pais).map(mapa_nomes_paises).fillna("Desconhecido")
                imp_p_c['País'] = imp_p_c['CO_PAIS'].apply(normalizar_pais).map(mapa_nomes_paises).fillna("Desconhecido")
                
                imp_final = pd.merge(imp_p_p, imp_p_c, on='País', how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                imp_final['Variação %'] = imp_final.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                imp_final = imp_final.sort_values(by=col_princ, ascending=False)
                imp_final = imp_final.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}'})
                
                df_show_i = imp_final.copy()
                df_show_i[f'Valor {ano_principal}'] = df_show_i[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_i[f'Valor {ano_comparacao}'] = df_show_i[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_i.head(top_n_itens), hide_index=True, use_container_width=True)

                st.subheader("Principais Produtos")
                imp_pr_p = df_imp_princ_f.groupby('SH4')['VL_FOB'].sum().reset_index()
                imp_pr_c = df_imp_comp_f.groupby('SH4')['VL_FOB'].sum().reset_index()
                
                imp_final_pr = pd.merge(imp_pr_p, imp_pr_c, on='SH4', how='outer', suffixes=(f' {ano_principal}', f' {ano_comparacao}')).fillna(0)
                imp_final_pr['Variação %'] = imp_final_pr.apply(lambda r: calc_var_display(r, col_princ, col_comp), axis=1)
                imp_final_pr = imp_final_pr.sort_values(by=col_princ, ascending=False)
                imp_final_pr['Descrição'] = imp_final_pr['SH4'].map(mapa_sh4_nomes).fillna("Desconhecido")
                
                imp_final_pr = imp_final_pr.rename(columns={col_princ: f'Valor {ano_principal}', col_comp: f'Valor {ano_comparacao}', 'SH4': 'Código SH4'})
                
                df_show_ip = imp_final_pr.copy()
                df_show_ip[f'Valor {ano_principal}'] = df_show_ip[f'Valor {ano_principal}'].apply(formatar_valor)
                df_show_ip[f'Valor {ano_comparacao}'] = df_show_ip[f'Valor {ano_comparacao}'].apply(formatar_valor)
                st.dataframe(df_show_ip[['Código SH4', 'Descrição', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_itens), hide_index=True, use_container_width=True)

                # Salvar
                file_bytes, file_name = app.finalizar_documento()
                st.session_state.arquivos_gerados_municipio.append({"name": file_name, "data": file_bytes})

        except Exception as e:
            st.error("Ocorreu um erro.")
            st.exception(e)

if st.session_state.arquivos_gerados_municipio:
    st.header("4. Relatórios Gerados")
    if len(st.session_state.arquivos_gerados_municipio) > 1:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for arquivo in st.session_state.arquivos_gerados_municipio:
                zip_file.writestr(arquivo["name"], arquivo["data"])
        st.download_button("Baixar ZIP", data=zip_buffer.getvalue(), file_name=f"Municipios_{ano_principal}.zip", mime="application/zip")
    else:
        arq = st.session_state.arquivos_gerados_municipio[0]
        st.download_button(f"Baixar {arq['name']}", data=arq["data"], file_name=arq['name'], mime="application/docx")