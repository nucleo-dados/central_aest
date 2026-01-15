import streamlit as st
import pandas as pd
import random
import os
import io
import requests
import unicodedata
import zipfile
from playwright.sync_api import sync_playwright

# --- IMPORTAÇÕES PARA DOCX ---
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 0. INTEGRAÇÃO E CONFIGURAÇÃO DA PÁGINA
# ==========================================

# 1. Configuração da aba do navegador
st.set_page_config(
    page_title="Briefings de Investimento",
    page_icon="💰",
    layout="wide"
)

# 2. Integração com auth.py (Segurança e Sidebar)
try:
    from auth import page_protector
except ImportError:
    st.error("Erro crítico: auth.py não encontrado. Execute a partir da raiz do projeto.")
    st.stop()

# 3. Aplica proteção e desenha sidebar personalizada
page_protector(page_name="Briefings de Investimento")

# 4. CSS para ocultar navegação padrão e ajustar layout
st.markdown(
    """
    <style>
        [data-testid="stSidebarNav"] { display: none; }
        .stButton button { width: 100%; }
        div[data-testid="stVerticalBlock"] > div { gap: 1rem; }
        .stAlert { margin-top: 1rem; }
    </style>
    """,
    unsafe_allow_html=True
)

# ==========================================
# 1. FUNÇÕES AUXILIARES E CLASSES
# ==========================================

def sanitize_filename(filename):
    import re
    return re.sub(r'[\\/*?:"<>|]', "_", str(filename))

@st.cache_data
def obter_municipios_ibge_mg():
    url = "https://servicodados.ibge.gov.br/api/v1/localidades/estados/MG/municipios"
    try:
        response = requests.get(url, timeout=5)
        response.raise_for_status()
        return {normalizar_string(i['nome']): i['nome'] for i in response.json()}
    except: return {}

@st.cache_data
def obter_paises_ibge():
    url = "https://servicodados.ibge.gov.br/api/v1/paises"
    try:
        response = requests.get(url, timeout=5)
        response.raise_for_status()
        return {normalizar_string(i['nome']['abreviado']): i['nome']['abreviado'] for i in response.json()}
    except: return {}

def normalizar_string(texto):
    if not isinstance(texto, str): return str(texto)
    nfkd = unicodedata.normalize('NFKD', texto)
    return "".join([c for c in nfkd if not unicodedata.combining(c)]).lower().strip()

def resolver_siglas_pais(nome):
    mapa = {'EUA': 'Estados Unidos', 'USA': 'Estados Unidos', 'UK': 'Reino Unido'}
    return mapa.get(str(nome).strip().upper(), nome)

def limpar_texto_proprio(texto):
    if pd.isna(texto) or str(texto).strip() in ["", "nan", "None"]: return "Não Informado"
    texto = str(texto).strip()
    preposicoes = ['de', 'da', 'do', 'dos', 'das', 'e', 'em', 'para']
    return " ".join([p.lower() if i>0 and p.lower() in preposicoes else (p.upper() if len(p)<5 and "." in p else p.capitalize()) for i, p in enumerate(texto.split())])

class DocumentoApp:
    def __init__(self, logo_path="LogoMinasGerais.png"):
        self.doc = Document()
        self.logo_path = logo_path
        self.configurar_estilos()

    def configurar_estilos(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        table = header.add_table(rows=1, cols=2, width=Cm(16.0))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(4.0)
        table.columns[1].width = Cm(12.0)
        
        c_img = table.cell(0, 0)
        p_img = c_img.paragraphs[0]
        run_img = p_img.add_run()
        if os.path.exists(self.logo_path):
            try: run_img.add_picture(self.logo_path, width=Cm(3.5), height=Cm(3.42))
            except: pass
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        c_txt = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        for i, t in enumerate(textos):
            p = c_txt.paragraphs[0] if i == 0 else c_txt.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)
            run = p.add_run(t)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = (i < 2)

    def set_titulo(self, titulo):
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(titulo.upper())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def adicionar_subtitulo(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(str(texto).upper())
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def adicionar_paragrafo(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_topico(self, texto_negrito, texto_normal):
        p = self.doc.add_paragraph(style='List Bullet')
        run_b = p.add_run(texto_negrito)
        run_b.font.name = 'Times New Roman'
        run_b.size = Pt(12)
        run_b.bold = True
        run_n = p.add_run(texto_normal)
        run_n.font.name = 'Times New Roman'
        run_n.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def finalizar_documento(self, nome_arquivo="Briefing.docx"):
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue(), sanitize_filename(nome_arquivo)

# ==========================================
# 2. CONFIGURAÇÕES E DADOS
# ==========================================

URL_SHAREPOINT = "https://indimg.sharepoint.com/sites/ArquivosPainelSEDE/Documentos%20Compartilhados/SEDE/EstoquePorData_SEDE.xlsx"
ARQUIVO_LOCAL = "base_estoque_temp.xlsx"
NOME_ABA = "EstoqueDataEstágio"

MAPA_COLUNAS = {
    'investimento': 'Investimentos', 'empresa': 'Empresa', 'projeto': 'Projeto',
    'empregos_dir': 'Empregos Diretos', 'empregos_ind': 'Empregos Indiretos',
    'municipio': 'Município', 'cadeia': 'Cadeia Produtiva',
    'situacao': 'Situação Projeto', 'territorio': 'Território de Desenvolvimento',
    'pais': 'País', 'data_df': 'DF'
}

def baixar_base_sharepoint():
    auth_file = "auth.json"
    if not os.path.exists(auth_file): return False, "Arquivo 'auth.json' ausente na raiz."
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(storage_state=auth_file, accept_downloads=True)
            page = context.new_page()
            url_dl = URL_SHAREPOINT + ("&download=1" if "?" in URL_SHAREPOINT else "?download=1")
            try:
                with page.expect_download(timeout=60000) as dl_info:
                    try: page.goto(url_dl)
                    except Exception as e:
                        if "Download is starting" not in str(e): raise e
                dl = dl_info.value
                dl.save_as(ARQUIVO_LOCAL)
                browser.close()
                return True, "Download ok."
            except Exception as e:
                browser.close()
                return False, f"Erro: {e}"
    except Exception as e: return False, f"Erro Crítico Playwright: {e}"

def formatar_moeda_humana(valor):
    try:
        val = float(valor)
        if val >= 1_000_000_000: return f"R$ {val/1_000_000_000:.2f} bilhões".replace('.', ',')
        elif val >= 1_000_000: return f"R$ {val/1_000_000:.2f} milhões".replace('.', ',')
        else: return f"R$ {val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except: return f"R$ {valor}"

def gerar_texto_empresa(dados):
    empresa = str(dados.get(MAPA_COLUNAS['empresa'], 'Empresa')).strip()
    projeto = str(dados.get(MAPA_COLUNAS['projeto'], 'Projeto')).strip()
    municipio = str(dados.get(MAPA_COLUNAS['municipio'], 'Local não informado')).strip()
    valor = formatar_moeda_humana(dados.get(MAPA_COLUNAS['investimento'], 0))
    
    try: n_dir = int(float(dados.get(MAPA_COLUNAS['empregos_dir'], 0)))
    except: n_dir = 0
    try: n_ind = int(float(dados.get(MAPA_COLUNAS['empregos_ind'], 0)))
    except: n_ind = 0

    total_vagas = n_dir + n_ind
    if total_vagas == 0: txt_emp = "sem estimativa de empregos"
    elif n_dir > 0 and n_ind > 0: txt_emp = f"gerando {total_vagas} empregos, sendo {n_dir} diretos e {n_ind} indiretos"
    elif n_dir > 0: txt_emp = f"gerando {n_dir} empregos diretos"
    else: txt_emp = f"gerando {n_ind} empregos indiretos"

    opcoes = [
        f": {projeto}. O investimento formalizado é de {valor}, com expectativa de {txt_emp}.",
        f": Em {municipio}, formalizou {valor}. Trata-se do {projeto}, {txt_emp}.",
        f": {projeto} ({municipio}). Aporte de {valor} e previsão de postos de trabalho: {txt_emp}.",
        f": {valor} formalizados em {municipio}. O foco é o {projeto}, {txt_emp}."
    ]
    return empresa, random.choice(opcoes)

# ==========================================
# 3. INTERFACE PRINCIPAL
# ==========================================
st.title("Briefings de Investimento (SEDE)")
st.markdown("---")

# --- BLOCO DE CARGA DE DADOS ---
col_load, _ = st.columns([1, 3])
with col_load:
    btn_load = st.button("🔄 Atualizar Base (SharePoint)", type="primary")

if 'df_raw' not in st.session_state: st.session_state.df_raw = None
if 'arquivos_gerados' not in st.session_state: st.session_state.arquivos_gerados = []

def limpar_cache_arquivos():
    st.session_state.arquivos_gerados = []

if btn_load:
    with st.status("Conectando ao SharePoint...", expanded=True) as status:
        ok, msg = baixar_base_sharepoint()
        if ok:
            try:
                st.write("Lendo arquivo Excel...")
                df = pd.read_excel(ARQUIVO_LOCAL, sheet_name=NOME_ABA)
                df.columns = df.columns.astype(str).str.strip()
                
                # Tratamentos
                st.write("Padronizando Municípios e Países (API IBGE)...")
                mapa_mun = obter_municipios_ibge_mg()
                mapa_pais = obter_paises_ibge()
                
                if MAPA_COLUNAS['municipio'] in df.columns:
                    df['chave'] = df[MAPA_COLUNAS['municipio']].astype(str).apply(normalizar_string)
                    df[MAPA_COLUNAS['municipio']] = df['chave'].map(mapa_mun).fillna(df[MAPA_COLUNAS['municipio']].apply(limpar_texto_proprio))
                
                col_pais = MAPA_COLUNAS['pais']
                if col_pais in df.columns:
                    df[col_pais] = df[col_pais].apply(resolver_siglas_pais)
                    df['chave_p'] = df[col_pais].astype(str).apply(normalizar_string)
                    df[col_pais] = df['chave_p'].map(mapa_pais).fillna(df[col_pais].apply(limpar_texto_proprio))

                for c in ['empresa', 'territorio', 'situacao', 'cadeia']:
                    if MAPA_COLUNAS[c] in df.columns: df[MAPA_COLUNAS[c]] = df[MAPA_COLUNAS[c]].apply(limpar_texto_proprio)

                if MAPA_COLUNAS['data_df'] in df.columns:
                    df['ano_temp'] = pd.to_datetime(df[MAPA_COLUNAS['data_df']], errors='coerce').dt.year
                    df['ano_filtro'] = df['ano_temp'].fillna(0).astype(int).astype(str).replace('0', 'Sem Data')
                else: df['ano_filtro'] = "N/A"

                st.session_state.df_raw = df
                limpar_cache_arquivos()
                status.update(label="Base Atualizada com Sucesso!", state="complete", expanded=False)
            except Exception as e:
                status.update(label="Erro no Processamento", state="error")
                st.error(f"Detalhe: {e}")
        else:
            status.update(label="Erro no Download", state="error")
            st.error(msg)

# SEÇÃO 1: CONFIGURAÇÕES
if st.session_state.df_raw is not None:
    df = st.session_state.df_raw
    
    st.header("1. Configurações da Análise")
    
    df_f = df.copy()
    c1, c2 = st.columns(2)
    
    with c1:
        opt_cad = sorted(df[MAPA_COLUNAS['cadeia']].unique())
        sel_cad = st.multiselect("Cadeia Produtiva:", opt_cad, on_change=limpar_cache_arquivos)
        if sel_cad: df_f = df_f[df_f[MAPA_COLUNAS['cadeia']].isin(sel_cad)]
        
        opt_mun = sorted(df_f[MAPA_COLUNAS['municipio']].unique())
        sel_mun = st.multiselect("Município:", opt_mun, on_change=limpar_cache_arquivos)
        if sel_mun: df_f = df_f[df_f[MAPA_COLUNAS['municipio']].isin(sel_mun)]
        
        if MAPA_COLUNAS['pais'] in df.columns:
            opt_pais = sorted(df_f[MAPA_COLUNAS['pais']].unique())
            sel_pais = st.multiselect("País:", opt_pais, on_change=limpar_cache_arquivos)
            if sel_pais: df_f = df_f[df_f[MAPA_COLUNAS['pais']].isin(sel_pais)]

        # Top N
        top_n = st.number_input("Top N Maiores Investimentos:", min_value=0, value=0, help="0 para trazer todos", on_change=limpar_cache_arquivos)

    with c2:
        opt_emp = sorted(df_f[MAPA_COLUNAS['empresa']].unique())
        sel_emp = st.multiselect("Empresa:", opt_emp, on_change=limpar_cache_arquivos)
        if sel_emp: df_f = df_f[df_f[MAPA_COLUNAS['empresa']].isin(sel_emp)]
        
        opt_sit = sorted(df_f[MAPA_COLUNAS['situacao']].unique())
        sel_sit = st.multiselect("Situação:", opt_sit, on_change=limpar_cache_arquivos)
        if sel_sit: df_f = df_f[df_f[MAPA_COLUNAS['situacao']].isin(sel_sit)]
        
        opt_ano = sorted(df_f['ano_filtro'].unique())
        sel_ano = st.multiselect("Ano de Referência:", opt_ano, on_change=limpar_cache_arquivos)
        if sel_ano: df_f = df_f[df_f['ano_filtro'].isin(sel_ano)]

    if top_n > 0:
        df_f[MAPA_COLUNAS['investimento']] = pd.to_numeric(df_f[MAPA_COLUNAS['investimento']], errors='coerce').fillna(0)
        df_f = df_f.sort_values(by=MAPA_COLUNAS['investimento'], ascending=False).head(top_n)

    st.info(f"Projetos selecionados para análise: **{len(df_f)}**")

    # SEÇÃO 2: AGRUPAMENTO
    st.header("2. Opções de Agrupamento")
    
    st.markdown("##### Foco Principal da Análise:")
    foco = st.radio("foco_analise", ["Cadeia Produtiva", "Município", "Território de Desenvolvimento"], horizontal=True, label_visibility="collapsed", on_change=limpar_cache_arquivos)
    
    col_foco = MAPA_COLUNAS['cadeia']
    selecao_foco = sel_cad
    if foco == "Município": 
        col_foco = MAPA_COLUNAS['municipio']
        selecao_foco = sel_mun
    elif foco == "Território de Desenvolvimento":
        col_foco = MAPA_COLUNAS['territorio']
        selecao_foco = df_f[col_foco].unique().tolist()
    
    if not selecao_foco: selecao_foco = df_f[col_foco].unique().tolist()
    
    agrupado = True
    tipo_geracao = "Agrupados"
    
    st.markdown(f"Deseja que os dados dos **{len(selecao_foco)} {foco}(s)** selecionados sejam agrupados?")
    tipo_geracao = st.radio("modo_grupo", ["Agrupados", "Separados"], horizontal=True, label_visibility="collapsed", on_change=limpar_cache_arquivos)
    
    st.info("""
    **💡 Como funciona o agrupamento:**
    * **Agrupados:** Gerará um **único relatório** consolidado. O texto somará os investimentos e empregos de todos os itens selecionados.
    * **Separados:** Gerará um **relatório individual** (em ZIP) para cada item do foco selecionado.
    """)

    nome_grupo = None
    if tipo_geracao == "Agrupados":
        agrupado = True
        st.markdown("**Dividir capítulos do documento por:**")
        divisao_caps = st.radio("div_caps", ["Cadeia Produtiva", "Município", "Território"], horizontal=True, label_visibility="collapsed")
        col_caps = MAPA_COLUNAS['cadeia']
        if divisao_caps == "Município": col_caps = MAPA_COLUNAS['municipio']
        elif divisao_caps == "Território": col_caps = MAPA_COLUNAS['territorio']
        
        if st.checkbox("Deseja dar um nome personalizado para este agrupamento?", value=True):
            nome_grupo = st.text_input("Digite o nome do agrupamento:", value=f"Briefing {foco}s Selecionados")
    else:
        agrupado = False
        col_caps = MAPA_COLUNAS['cadeia'] 

    # SEÇÃO 3: GERAR
    st.header("3. Gerar Análise")
    
    if st.button("Iniciar Análise de Investimentos", type="primary"):
        st.session_state.arquivos_gerados = []
        
        with st.spinner("Gerando documentos..."):
            items_processar = []
            
            if agrupado:
                titulo = nome_grupo if nome_grupo else "Briefing de Investimentos"
                items_processar.append({ "titulo": titulo, "df": df_f, "col_cap": col_caps })
            else:
                for item in selecao_foco:
                    d_sub = df_f[df_f[col_foco] == item]
                    if len(d_sub) > 0:
                        items_processar.append({ "titulo": f"Briefing - {item}", "df": d_sub, "col_cap": col_caps })
            
            for task in items_processar:
                df_t = task['df']
                tit = task['titulo']
                cc = task['col_cap']
                
                # Prepara números
                for c in [MAPA_COLUNAS['investimento'], MAPA_COLUNAS['empregos_dir'], MAPA_COLUNAS['empregos_ind']]:
                    df_t[c] = pd.to_numeric(df_t[c], errors='coerce').fillna(0)
                
                sum_inv = formatar_moeda_humana(df_t[MAPA_COLUNAS['investimento']].sum())
                sum_dir = int(df_t[MAPA_COLUNAS['empregos_dir']].sum())
                sum_ind = int(df_t[MAPA_COLUNAS['empregos_ind']].sum())
                sum_tot = sum_dir + sum_ind
                
                # Texto Intro Ajustado
                if sum_dir > 0 and sum_ind > 0: txt_intro_emp = f"com previsão de {sum_tot} empregos, sendo {sum_dir} diretos e {sum_ind} indiretos"
                elif sum_dir > 0: txt_intro_emp = f"com previsão de {sum_dir} empregos diretos"
                elif sum_ind > 0: txt_intro_emp = f"com previsão de {sum_ind} empregos indiretos"
                else: txt_intro_emp = "sem estimativa de geração de empregos"
                
                # Gera Word
                doc = DocumentoApp()
                doc.set_titulo(tit)
                doc.adicionar_paragrafo(f"O recorte analisado (2019-2025) totaliza investimentos de {sum_inv} em {len(df_t)} projetos, {txt_intro_emp}.")
                
                df_t[cc] = df_t[cc].fillna("Outros")
                df_t = df_t.sort_values(by=MAPA_COLUNAS['investimento'], ascending=False)
                
                for g, sub in df_t.groupby(cc):
                    doc.adicionar_subtitulo(str(g))
                    for reg in sub.to_dict('records'):
                        e, t = gerar_texto_empresa(reg)
                        doc.adicionar_topico(e, t)
                
                b, n = doc.finalizar_documento(f"{tit}.docx")
                st.session_state.arquivos_gerados.append({"name": n, "data": b})

# DOWNLOAD
if st.session_state.arquivos_gerados:
    st.divider()
    st.markdown("### 📥 Relatórios Gerados")
    arquivos = st.session_state.arquivos_gerados
    
    if len(arquivos) == 1:
        st.success("Relatório pronto.")
        st.download_button(f"Baixar {arquivos[0]['name']}", data=arquivos[0]['data'], file_name=arquivos[0]['name'], mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.success(f"{len(arquivos)} relatórios gerados.")
        zb = io.BytesIO()
        with zipfile.ZipFile(zb, "w", zipfile.ZIP_DEFLATED) as zf:
            for a in arquivos: zf.writestr(a['name'], a['data'])
        st.download_button(f"Baixar Pacote ZIP ({len(arquivos)} arquivos)", data=zb.getvalue(), file_name="Briefings_Investimentos.zip", mime="application/zip")
else:
    if st.session_state.df_raw is not None:
        st.caption("Configure as opções acima e clique no botão para gerar.")
