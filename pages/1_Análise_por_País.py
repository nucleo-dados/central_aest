import streamlit as st
import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime
import io
import re
import zipfile
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

# --- IMPORTAÇÃO E PROTEÇÃO DA PÁGINA ---
try:
    from auth import page_protector 
    # Protege a página, oculta nav padrão e desenha a sidebar
    page_protector(page_name="Análise por País")
except ImportError:
    # Caso esteja rodando localmente sem o auth.py, apenas avisa mas não para (para testes)
    st.warning("Atenção: Módulo de autenticação 'auth' não encontrado. Rodando em modo de teste.")

# --- CONFIGURAÇÕES GLOBAIS E CONSTANTES ---
requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)
estados_brasileiros = {'AC', 'AL', 'AP', 'AM', 'BA', 'CE', 'DF', 'ES', 'GO', 'MA', 'MT', 'MS', 'MG', 'PA', 'PB', 'PR',
                       'PE', 'PI', 'RJ', 'RN', 'RS', 'RO', 'RR', 'SC', 'SE', 'SP', 'TO'}
meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}
MESES_MAPA = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
    "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}
LISTA_MESES = list(MESES_MAPA.keys())
ARTIGOS_PAISES_MAP = {
    "Afeganistão": "o", "África do Sul": "a", "Alemanha": "a", "Arábia Saudita": "a",
    "Argentina": "a", "Austrália": "a", "Bélgica": "a", "Brasil": "o", "Canadá": "o",
    "Chade": "o", "Chile": "o", "China": "a", "Colômbia": "a", "Congo": "o",
    "Coreia do Norte": "a", "Coreia do Sul": "a", "Costa Rica": "a", "Equador": "o",
    "Egito": "o", "Emirados Árabes Unidos": "os", "Espanha": "a", "Estados Unidos": "os",
    "Filipinas": "as", "França": "a", "Holanda": "a", "Índia": "a", "Indonésia": "a",
    "Inglaterra": "a", "Irã": "o", "Itália": "a", "Japão": "o", "Líbano": "o",
    "Malásia": "a", "México": "o", "Nicarágua": "a", "Noruega": "a", "Nova Zelândia": "a",
    "Países Baixos": "os", "Panamá": "o", "Paraguai": "o", "Pérsia": "a", "Peru": "o",
    "Reino Unido": "o", "República Checa": "a", "República Dominicana": "a",
    "Romênia": "a", "Rússia": "a", "Singapura": "a", "Suécia": "a", "Uruguai": "o",
    "Venezuela": "a", "Vietnã": "o"
}

# --- BLOCO MANUAL DE BLOCOS ECONÔMICOS ---
BLOCOS_ECONOMICOS = {
    "América Central e Caribe": [],
    "América do Norte": ["Estados Unidos", "Canadá", "México"],
    "América do Sul": ["Argentina", "Bolívia", "Chile", "Colômbia", "Equador", "Guiana", "Paraguai", "Peru", "Suriname", "Uruguai", "Venezuela"],
    "Associação de Nações do Sudeste Asiático - ASEAN": ["Brunei", "Camboja", "Filipinas", "Indonésia", "Laos", "Malásia", "Myanmar", "Singapura", "Tailândia", "Vietnã"],
    "Comunidade Andina das Nações - CAN": ["Bolívia", "Colômbia", "Equador", "Peru"],
    "Europa": ["Albânia", "Alemanha", "Andorra", "Áustria", "Bélgica", "Bielorrússia", "Bósnia-Herzegovina", "Bulgária", "Chipre", "Croácia", "Dinamarca", "Eslováquia", "Eslovênia", "Espanha", "Estônia", "Finlândia", "França", "Grécia", "Hungria", "Irlanda", "Islândia", "Itália", "Letônia", "Listenstaine", "Lituânia", "Luxemburgo", "Macedônia do Norte", "Malta", "Moldávia", "Mônaco", "Montenegro", "Noruega", "Países Baixos", "Polônia", "Portugal", "Reino Unido", "República Checa", "Romênia", "Rússia", "San Marino", "Sérvia", "Suécia", "Suíça", "Ucrânia", "Vaticano"],
    "Mercado Comum do Sul - Mercosul": ["Argentina", "Paraguai", "Uruguai", "Venezuela"],
    "Oceania": ["Austrália", "Fiji", "Ilhas Marshall", "Ilhas Salomão", "Kiribati", "Micronésia", "Nauru", "Nova Zelândia", "Palau", "Papua Nova Guiné", "Samoa", "Tonga", "Tuvalu", "Vanuatu"],
    "Oriente Médio": ["Arábia Saudita", "Bahrein", "Catar", "Emirados Árabes Unidos", "Iêmen", "Irã", "Iraque", "Israel", "Jordânia", "Kuwait", "Líbano", "Omã", "Palestina", "Síria", "Turquia"],
    "União Europeia - UE": ["Alemanha", "Áustria", "Bélgica", "Bulgária", "Chipre", "Croácia", "Dinamarca", "Eslováquia", "Eslovênia", "Espanha", "Estônia", "Finlândia", "França", "Grécia", "Hungria", "Irlanda", "Itália", "Letônia", "Lituânia", "Luxemburgo", "Malta", "Países Baixos", "Polônia", "Portugal", "República Checa", "Romênia", "Suécia"],
    "África": ["África do Sul", "Angola", "Argélia", "Benin", "Botsuana", "Burkina Faso", "Burundi", "Cabo Verde", "Camarões", "Chade", "Comores", "Congo", "Costa do Marfim", "Djibuti", "Egito", "Eritreia", "Eswatini", "Etiópia", "Gabão", "Gâmbia", "Gana", "Guiné", "Guiné Equatorial", "Guiné-Bissau", "Lesoto", "Libéria", "Líbia", "Madagascar", "Malawi", "Mali", "Marrocos", "Maurício", "Mauritânia", "Moçambique", "Namíbia", "Níger", "Nigéria", "Quênia", "República Centro-Africana", "República Democrática do Congo", "Ruanda", "São Tomé e Príncipe", "Senegal", "Seychelles", "Serra Leoa", "Somália", "Sudão", "Sudão do Sul", "Tanzânia", "Togo", "Tunísia", "Uganda", "Zâmbia", "Zimbábue"],
    "Ásia (Exclusive Oriente Médio)": ["Afeganistão", "Armênia", "Azerbaijão", "Bangladesh", "Brunei", "Butão", "Camboja", "Cazaquistão", "China", "Coreia do Norte", "Coreia do Sul", "Filipinas", "Geórgia", "Índia", "Indonésia", "Japão", "Laos", "Malásia", "Maldivas", "Mongólia", "Myanmar", "Nepal", "Paquistão", "Quirguistão", "Singapura", "Sri Lanka", "Tailândia", "Tajiquistão", "Timor-Leste", "Turcomenistão", "Uzbequistão", "Vietnã"]
}
# --- FIM DO BLOCO MANUAL ---

NCM_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_NCM', 'CO_NCM']
NCM_DTYPES = {'CO_NCM': str, 'CO_SH4': str} 
MUN_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_MUN', 'CO_MUN']
MUN_DTYPES = {'CO_MUN': str}

# --- FUNÇÕES DE LÓGICA (Helpers) ---

@st.cache_data(ttl=3600)
def ler_dados_csv_online(url, usecols=None, dtypes=None):
    retries = 3
    for attempt in range(retries):
        try:
            resposta = requests.get(url, verify=False, timeout=(10, 1200)) 
            resposta.raise_for_status()
            final_dtypes = {'CO_SH4': str, 'CO_NCM': str}
            if dtypes:
                final_dtypes.update(dtypes)
            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1',
                             sep=';', 
                             dtype=final_dtypes,
                             usecols=usecols)
            return df
        except requests.exceptions.RequestException as e:
            print(f"Erro ao acessar o CSV (tentativa {attempt + 1}/{retries}): {e}")
            if "Read timed out" in str(e) and attempt < retries - 1:
                st.warning("Download demorou muito. Tentando novamente...")
                continue
            if "IncompleteRead" in str(e) and attempt < retries - 1:
                st.warning("Retentando download...")
                continue
            else:
                return None
        except Exception as e:
            print(f"Erro inesperado ao baixar ou processar o CSV: {e}")
            return None
    return None

@st.cache_data(ttl=3600)
def carregar_dataframe(url, nome_arquivo, usecols=None, dtypes=None, mostrar_progresso=True):
    progress_bar = None
    if mostrar_progresso: 
        progress_bar = st.progress(0, text=f"Carregando {nome_arquivo}...")
    
    df = ler_dados_csv_online(url, usecols=usecols, dtypes=dtypes)
    
    if mostrar_progresso and progress_bar: 
        if df is not None:
            progress_bar.progress(100, text=f"{nome_arquivo} carregado com sucesso.")
        else:
            progress_bar.empty()
    return df

@st.cache_data
def obter_dados_paises():
    """Carrega a tabela de países (ID e Nome) e armazena em cache."""
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv", usecols=['NO_PAIS', 'CO_PAIS'], mostrar_progresso=False) 
    if df_pais is not None and not df_pais.empty:
        mapa_codigo_nome = pd.Series(df_pais.NO_PAIS.values, index=df_pais.CO_PAIS).to_dict()
        lista_nomes = sorted(df_pais[df_pais['NO_PAIS'] != 'Brasil']['NO_PAIS'].unique().tolist())
        mapa_nome_codigo = pd.Series(df_pais.CO_PAIS.values, index=df_pais.NO_PAIS).to_dict()
        return mapa_codigo_nome, lista_nomes, mapa_nome_codigo
    return {}, [], {}

@st.cache_data
def obter_dados_produtos_ncm():
    """Carrega a tabela NCM (SH4) e armazena em cache."""
    url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
    usecols_ncm = ['CO_SH4', 'NO_SH4_POR']
    df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv", usecols=usecols_ncm, dtypes={'CO_SH4': str}, mostrar_progresso=False)
    if df_ncm is not None:
        df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].astype(str).str.zfill(4)
        mapa_sh4 = df_ncm.drop_duplicates('CO_SH4_STR').set_index('CO_SH4_STR')['NO_SH4_POR']
        return df_ncm, mapa_sh4.to_dict()
    return None, {}

def get_sh4(co_ncm):
    """Extrai SH4 de um CO_NCM."""
    co_ncm_str = str(co_ncm).strip()
    if pd.isna(co_ncm_str) or co_ncm_str == "":
        return None
    co_ncm_str = co_ncm_str.zfill(8)
    return co_ncm_str[:4]

@st.cache_data
def obter_lista_de_blocos():
    """Retorna uma lista de nomes de blocos econômicos (hardcoded)."""
    blocos = sorted(list(BLOCOS_ECONOMICOS.keys()))
    return blocos

@st.cache_data
def obter_paises_do_bloco(nome_bloco):
    """Retorna uma lista de nomes de países (hardcoded) para um bloco específico."""
    return BLOCOS_ECONOMICOS.get(nome_bloco, [])

def obter_lista_de_paises(lista_nomes):
    """Apenas retorna a lista de nomes pré-carregada."""
    if not lista_nomes:
        return ["Erro ao carregar lista de países"]
    return lista_nomes

def obter_codigo_pais(nome_pais, mapa_reverso):
    """Obtém o código do país a partir do mapa."""
    return mapa_reverso.get(nome_pais)

def validar_paises(paises_selecionados, mapa_nome_codigo):
    """Valida a lista de países usando o mapa pré-carregado."""
    codigos_paises = []
    nomes_paises_validos = []
    paises_invalidos = []
    
    for pais in paises_selecionados:
        if pais.lower() == "brasil":
            paises_invalidos.append(f"{pais} (Não é possível fazer busca no Brasil)")
            continue
        codigo_pais = mapa_nome_codigo.get(pais) 
        if codigo_pais is None:
            paises_invalidos.append(f"{pais} (País não encontrado)")
        else:
            codigos_paises.append(codigo_pais)
            nomes_paises_validos.append(pais)
    return codigos_paises, nomes_paises_validos, paises_invalidos

def filtrar_dados_por_estado_e_mes(df, estados, meses_para_filtrar):
    df_filtrado = df[df['SG_UF_NCM'].isin(list(estados))]
    df_filtrado = df_filtrado[df_filtrado['CO_MES'].isin(meses_para_filtrar)]
    return df_filtrado

def filtrar_dados_por_mg_e_pais(df, codigos_paises, agrupado, meses_para_filtrar):
    df_filtrado = df[df['SG_UF_NCM'] == 'MG']
    if agrupado:
        df_filtrado = df_filtrado[df_filtrado['CO_PAIS'].isin(codigos_paises)]
    else:
        # Se não agrupado, assume lista de 1 item
        if isinstance(codigos_paises, list) and len(codigos_paises) > 0:
             df_filtrado = df_filtrado[df_filtrado['CO_PAIS'] == codigos_paises[0]]
        else:
             df_filtrado = df_filtrado[df_filtrado['CO_PAIS'] == codigos_paises]

    df_filtrado = df_filtrado[df_filtrado['CO_MES'].isin(meses_para_filtrar)]
    return df_filtrado

def calcular_ranking_por_pais(df):
    ranking = df.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
    return ranking

def calcular_participacao(valor_parcial, valor_total):
    if valor_total == 0:
        return 0.0
    participacao = round(valor_parcial / valor_total * 100, 2)
    return participacao

def calcular_diferenca_percentual(valor_atual, valor_anterior):
    if valor_anterior == 0:
        return 0.0, "acréscimo" if valor_atual > 0 else "redução" if valor_atual < 0 else "estabilidade"
    diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    if diferenca > 0:
        tipo_diferenca = "um acréscimo"
    elif diferenca < 0:
        tipo_diferenca = "uma queda" # Alterado para bater com o doc enviado ("queda" ao inves de "redução")
    else:
        tipo_diferenca = "uma estabilidade"
    diferenca = abs(diferenca)
    return diferenca, tipo_diferenca

def calcular_posicao_estado_pais(df, codigos_paises):
    """
    Calcula a posição de MG no ranking dos estados (ou UFs) que exportam/importam 
    para o conjunto de países selecionados.
    """
    # Filtra transações apenas para os países de interesse
    df_filtrado = df[df['CO_PAIS'].isin(codigos_paises)]
    
    if df_filtrado.empty:
        return 0
    
    # Agrupa por UF e soma o valor FOB
    ranking_uf = df_filtrado.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(ascending=False)
    
    # Verifica se MG está no ranking
    if 'MG' not in ranking_uf.index:
        return 0
        
    # Retorna a posição (índice 0-based + 1)
    return ranking_uf.index.get_loc('MG') + 1

def calcular_ranking_e_participacao_brasil(df_brasil, codigos_paises):
    """
    Calcula o ranking de MG entre os estados brasileiros para um destino/origem
    e a participação de MG no total do Brasil.
    """
    # Filtra dados do Brasil para o país/bloco destino
    df_brasil_pais = df_brasil[df_brasil['CO_PAIS'].isin(codigos_paises)]
    
    # Agrupa por UF
    ranking_uf = df_brasil_pais.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(ascending=False)
    
    total_brasil_pais = ranking_uf.sum()
    
    if 'MG' not in ranking_uf.index:
        return 0, 0.0
        
    posicao_mg = ranking_uf.index.get_loc('MG') + 1
    valor_mg = ranking_uf['MG']
    
    participacao_mg_br = 0.0
    if total_brasil_pais > 0:
        participacao_mg_br = round((valor_mg / total_brasil_pais) * 100, 2)
        
    return posicao_mg, participacao_mg_br

def calcular_balanca_e_fluxo(exportacao_ano, importacao_ano, exportacao_ano_anterior, importacao_ano_anterior):
    balanca_ano = exportacao_ano - importacao_ano
    balanca_ano_anterior = exportacao_ano_anterior - importacao_ano_anterior
    fluxo_comercial_ano = exportacao_ano + importacao_ano
    fluxo_comercial_ano_anterior = exportacao_ano_anterior + importacao_ano_anterior
    variacao_balanca = 0
    variacao_fluxo = 0
    if balanca_ano_anterior != 0:
        variacao_balanca = ((balanca_ano - balanca_ano_anterior) / balanca_ano_anterior) * 100
    if fluxo_comercial_ano_anterior != 0:
        variacao_fluxo = ((fluxo_comercial_ano - fluxo_comercial_ano_anterior) / fluxo_comercial_ano_anterior) * 100
    return balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo

def gerar_texto_lista_produtos(df_dados, mapa_nomes, top_n=5):
    """Gera string: 'Produto A (X%); Produto B (Y%); ...' """
    if df_dados.empty:
        return "Nenhum produto registrado."
        
    total = df_dados['VL_FOB'].sum()
    if total == 0:
        return "Valor total zero."
        
    agrupado = df_dados.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(top_n)
    lista_textos = []
    
    for sh4, valor in agrupado.items():
        nome = mapa_nomes.get(sh4, "Produto Desconhecido")
        part = (valor / total) * 100
        lista_textos.append(f"{nome} ({part:.2f}%)")
        
    return "; ".join(lista_textos) + "."

def gerar_texto_lista_municipios(df_dados, df_uf_mun, top_n=5):
    """Gera string: 'Município A (X%); Município B (Y%); ...' e retorna contagem total."""
    if df_dados.empty:
        return "Nenhum município.", 0
        
    total = df_dados['VL_FOB'].sum()
    if total == 0:
        return "Valor total zero.", 0

    # Mapa de códigos para nomes
    mapa_mun = pd.Series(df_uf_mun.NO_MUN_MIN.values, index=df_uf_mun.CO_MUN_GEO.astype(str)).to_dict()

    # Ajuste CO_MUN se necessário (as vezes vem como int ou str)
    df_dados = df_dados.copy()
    df_dados['CO_MUN'] = df_dados['CO_MUN'].astype(str)
    
    contagem_total = df_dados['CO_MUN'].nunique()
    
    agrupado = df_dados.groupby('CO_MUN')['VL_FOB'].sum().sort_values(ascending=False).head(top_n)
    lista_textos = []
    
    for co_mun, valor in agrupado.items():
        nome = mapa_mun.get(co_mun, f"Município {co_mun}")
        part = (valor / total) * 100
        lista_textos.append(f"{nome} ({part:.2f}%)")
        
    return "; ".join(lista_textos) + ".", contagem_total

def obter_artigo_pais(nome_pais):
    return ARTIGOS_PAISES_MAP.get(nome_pais, "") 

def formatar_valor(valor):
    prefixo = ""
    if valor < 0:
        prefixo = "-"
        valor = abs(valor)
    if valor >= 1_000_000_000:
        valor_formatado_str = f"{(valor / 1_000_000_000):.2f}".replace('.',',')
        unidade = "bilhão" if (valor / 1_000_000_000) < 2 else "bilhões"
        return f"{prefixo}US$ {valor_formatado_str} {unidade}"
    if valor >= 1_000_000:
        valor_formatado_str = f"{(valor / 1_000_000):.2f}".replace('.',',')
        unidade = "milhão" if (valor / 1_000_000) < 2 else "milhões"
        return f"{prefixo}US$ {valor_formatado_str} {unidade}"
    if valor >= 1_000:
        valor_formatado_str = f"{(valor / 1_000):.2f}".replace('.',',')
        return f"{prefixo}US$ {valor_formatado_str} mil"
    valor_formatado_str = f"{valor:.2f}".replace('.',',')
    return f"{prefixo}US$ {valor_formatado_str}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

class DocumentoApp:
    def __init__(self, logo_path):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""
        self.logo_path = logo_path
        self.diretorio_base = "/tmp/" 

    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Espaço após o título
        self.doc.add_paragraph()

    def adicionar_paragrafo(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def adicionar_titulo(self, texto):
        # Espaço antes do título da seção
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        if self.subsecao_atual == 0:
            run = p.add_run(f"{self.secao_atual}. {texto}")
        else:
            run = p.add_run(f"{self.secao_atual}.{self.subsecao_atual}. {texto}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def nova_secao(self):
        self.secao_atual += 1
        self.subsecao_atual = 0

    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        largura_total_cm = 16.0
        table = header.add_table(rows=1, cols=2, width=Cm(largura_total_cm))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(4.0)
        table.columns[1].width = Cm(12.0)
        cell_imagem = table.cell(0, 0)
        paragraph_imagem = cell_imagem.paragraphs[0]
        paragraph_imagem.paragraph_format.space_before = Pt(0)
        paragraph_imagem.paragraph_format.space_after = Pt(0)
        run_imagem = paragraph_imagem.add_run()
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                run_imagem.add_picture(self.logo_path,
                                       width=Cm(3.5), 
                                       height=Cm(3.42))
            except Exception as e:
                paragraph_imagem.add_run("[Logo não encontrado]")
        else:
            paragraph_imagem.add_run("")
        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        def formatar_paragrafo_cabecalho(p):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p = cell_texto.paragraphs[0]
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[0])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True 
        p = cell_texto.add_paragraph()
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[1])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        for texto in textos[2:]: 
            p = cell_texto.add_paragraph()
            formatar_paragrafo_cabecalho(p)
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = False 

    def finalizar_documento(self):
        diretorio_real = self.diretorio_base
        try:
            os.makedirs(diretorio_real, exist_ok=True)
        except Exception:
            diretorio_real = "/tmp/"
            os.makedirs(diretorio_real, exist_ok=True)
        nome_arquivo = f"{self.titulo_doc}.docx"
        nome_arquivo_sanitizado = sanitize_filename(nome_arquivo)
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        file_bytes = file_stream.getvalue()
        # Não salvamos no disco do servidor para evitar lotação, retornamos os bytes
        return file_bytes, nome_arquivo_sanitizado

# --- ----------------------------------- ---
# --- INTERFACE GRÁFICA DO STREAMLIT (Página 1) ---
# --- ----------------------------------- ---

# --- Inicialização do Session State ---
if 'arquivos_gerados_pais' not in st.session_state:
    st.session_state.arquivos_gerados_pais = []

# --- Callback para limpar o state ---
def clear_download_state_pais():
    if 'arquivos_gerados_pais' in st.session_state:
        st.session_state.arquivos_gerados_pais = []

# --- ENTRADAS PRINCIPAIS ---
st.header("1. Configurações da Análise")

try:
    mapa_nomes_paises, lista_paises_nomes, mapa_paises_reverso = obter_dados_paises()
    lista_de_blocos = obter_lista_de_blocos()
    df_ncm_completo, mapa_sh4_nomes = obter_dados_produtos_ncm() 
except Exception as e:
    st.error(f"Erro crítico ao carregar listas iniciais: {e}")
    lista_paises_nomes = ["Falha ao carregar países"]
    lista_de_blocos = ["Falha ao carregar blocos"]
    mapa_nomes_paises = {}
    mapa_paises_reverso = {}
    mapa_sh4_nomes = {}
    df_ncm_completo = None

lista_de_paises = obter_lista_de_paises(lista_paises_nomes)

# --- Lógica de 'default' resiliente ---
valores_padrao = ["China", "Estados Unidos"]
valores_padrao_filtrados = [pais for pais in valores_padrao if pais in lista_de_paises]
if not valores_padrao_filtrados and len(lista_de_paises) > 0 and "Erro" not in lista_de_paises[0]:
    valores_padrao_filtrados = [lista_de_paises[0]]
elif "Erro" in lista_de_paises[0] or "Falha" in lista_de_paises[0]:
    valores_padrao_filtrados = [] 
    st.warning("Não foi possível carregar a lista de países. O site de dados pode estar fora do ar.")
# --- FIM DA LÓGICA ---

ano_atual = datetime.now().year

col1, col2 = st.columns(2)
with col1:
    ano_principal = st.number_input(
        "Ano de Referência:",
        min_value=1998,
        max_value=ano_atual,
        value=ano_atual,
        help="O ano principal que você quer analisar.",
        on_change=clear_download_state_pais
    )
    ano_comparacao = st.number_input(
        "Ano de Comparação:",
        min_value=1998,
        max_value=ano_atual,
        value=ano_atual - 1,
        help="O ano contra o qual você quer comparar.",
        on_change=clear_download_state_pais
    )
    meses_selecionados = st.multiselect(
        "Meses de Análise (opcional):",
        options=LISTA_MESES,
        help="Selecione os meses. Se deixar em branco, o ano inteiro será analisado.",
        on_change=clear_download_state_pais
    )

with col2:
    blocos_selecionados = st.multiselect(
        "Filtrar por Bloco(s) (opcional):",
        options=lista_de_blocos,
        help="Os países destes blocos serão adicionados à seleção.",
        on_change=clear_download_state_pais
    )
    
    paises_selecionados_manual = st.multiselect(
        "Filtrar por País(es) (opcional):",
        options=lista_de_paises,
        default=valores_padrao_filtrados,
        help="Você pode digitar para pesquisar e selecionar múltiplos países.",
        on_change=clear_download_state_pais
    )
    
    top_n_produtos = st.number_input(
        "Nº de Produtos no Ranking:",
        min_value=1,
        max_value=100,
        value=10,
        help="Quantos produtos (SH4) devem ser exibidos nas tabelas de ranking e no texto.",
        on_change=clear_download_state_pais
    )

# --- LÓGICA CONDICIONAL PARA ENTRADAS ---
agrupado = True 
nome_agrupamento = None

paises_do_bloco = []
if blocos_selecionados:
    for bloco in blocos_selecionados:
        paises_do_bloco.extend(obter_paises_do_bloco(bloco))

paises = sorted(list(set(paises_selecionados_manual + paises_do_bloco)))

if len(paises) > 1:
    st.header("2. Opções de Agrupamento")
    
    if blocos_selecionados and not paises_selecionados_manual:
        agrupado = True
        st.info(f"Análise de Bloco será agrupada.")
        nome_agrupamento = ", ".join(blocos_selecionados)
    else:
        agrupamento_input = st.radio(
            "Deseja que os dados sejam agrupados ou separados?",
            ("agrupados", "separados"),
            index=0,
            horizontal=True,
            on_change=clear_download_state_pais
        )
        agrupado = (agrupamento_input == "agrupados")

        if agrupado:
            st.info(
                "💡 **Como funciona o agrupamento:**\n"
                "* **Agrupados:** Gerará um **único relatório** consolidado.\n"
                "* **Separados:** Gerará um **relatório individual** para cada país. O download será um arquivo .zip."
            )
            
            quer_nome_agrupamento = st.checkbox(
                "Deseja dar um nome para este agrupamento?", 
                key="pais_nome_grupo",
                on_change=clear_download_state_pais
            )
            if quer_nome_agrupamento:
                nome_agrupamento = st.text_input(
                    "Digite o nome do agrupamento:", 
                    key="pais_nome_input",
                    on_change=clear_download_state_pais
                )
    st.header("3. Gerar Relatório")
else:
    agrupado = False 
    st.header("2. Gerar Relatório")


# --- EXECUÇÃO DO SCRIPT ---
if st.button(" Iniciar Geração do Relatório"):
    
    st.session_state.arquivos_gerados_pais = []
    
    logo_path_to_use = "LogoMinasGerais.png" 
    if not os.path.exists(logo_path_to_use):
        st.warning(f"Aviso: A logo 'LogoMinasGerais.png' não foi encontrada. O cabeçalho será gerado sem a logo.")
        logo_path_to_use = None
    
    with st.spinner(f"Gerando relatório para {', '.join(paises)} ({ano_principal} vs {ano_comparacao})... Isso pode levar alguns minutos."):
        
        try:
            codigos_paises, nomes_paises_validos, paises_invalidos = validar_paises(paises, mapa_paises_reverso)
            if paises_invalidos:
                st.warning(f"Países não encontrados ou inválidos (ignorados): {', '.join(paises_invalidos)}")
            if not nomes_paises_validos:
                st.error("Nenhum país válido fornecido. A geração foi interrompida.")
                st.stop()
            
            url_exp_ano_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_principal}.csv"
            url_exp_ano_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_comparacao}.csv"
            url_imp_ano_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_principal}.csv"
            url_imp_ano_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_comparacao}.csv"
            url_exp_mun_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_principal}_MUN.csv"
            url_imp_mun_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_principal}_MUN.csv"
            url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
            
            df_ncm = df_ncm_completo 
            df_uf_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['CO_MUN_GEO', 'NO_MUN_MIN'], dtypes={'CO_MUN_GEO': str}, mostrar_progresso=False)
            
            if df_ncm is None or df_uf_mun is None:
                st.error("Não foi possível carregar tabelas auxiliares (NCM ou Municípios). Abortando.")
                st.stop()

            df_exp_ano = carregar_dataframe(url_exp_ano_principal, f"EXP_{ano_principal}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            df_exp_ano_anterior = carregar_dataframe(url_exp_ano_comparacao, f"EXP_{ano_comparacao}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)

            if df_exp_ano is None or df_exp_ano_anterior is None:
                st.error("Não foi possível carregar dados de exportação. Verifique os anos selecionados ou tente novamente mais tarde.")
                st.stop()
                
            df_exp_ano['SH4'] = df_exp_ano['CO_NCM'].apply(get_sh4)
            df_exp_ano_anterior['SH4'] = df_exp_ano_anterior['CO_NCM'].apply(get_sh4)

            ultimo_mes_disponivel = df_exp_ano['CO_MES'].max()
            meses_para_filtrar = []
            
            if not meses_selecionados: 
                meses_para_filtrar = list(range(1, ultimo_mes_disponivel + 1))
                nome_periodo = f"o ano de {ano_principal}"
                nome_periodo_em = f"Em {ano_principal}"
                nome_periodo_comp = f"{ano_comparacao}"
            else:
                meses_para_filtrar = [MESES_MAPA[m] for m in meses_selecionados]
                if max(meses_para_filtrar) > ultimo_mes_disponivel:
                    st.error(f"O ano {ano_principal} só possui dados até {meses_pt[ultimo_mes_disponivel]}. Por favor, desmarque os meses posteriores.")
                    st.stop()
                nome_periodo = f"o período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_em = f"No período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            
            # Filtros Gerais
            df_exp_ano_estados = filtrar_dados_por_estado_e_mes(df_exp_ano, estados_brasileiros, meses_para_filtrar)
            df_exp_ano_mg = filtrar_dados_por_estado_e_mes(df_exp_ano, ['MG'], meses_para_filtrar)
            
            # IMPORTAÇÕES
            df_imp_ano = carregar_dataframe(url_imp_ano_principal, f"IMP_{ano_principal}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            df_imp_ano_anterior = carregar_dataframe(url_imp_ano_comparacao, f"IMP_{ano_comparacao}.csv", usecols=NCM_COLS, dtypes=NCM_DTYPES)
            
            if df_imp_ano is None or df_imp_ano_anterior is None:
                st.error("Não foi possível carregar dados de importação. Abortando.")
                st.stop()
            
            df_imp_ano['SH4'] = df_imp_ano['CO_NCM'].apply(get_sh4)
            df_imp_ano_anterior['SH4'] = df_imp_ano_anterior['CO_NCM'].apply(get_sh4)
            
            df_imp_ano_estados = filtrar_dados_por_estado_e_mes(df_imp_ano, estados_brasileiros, meses_para_filtrar)
            df_imp_ano_mg = filtrar_dados_por_estado_e_mes(df_imp_ano, ['MG'], meses_para_filtrar)
            
            # DFs Municipais
            df_exp_mun = carregar_dataframe(url_exp_mun_principal, f"EXP_{ano_principal}_MUN.csv", usecols=MUN_COLS, dtypes=MUN_DTYPES)
            df_imp_mun = carregar_dataframe(url_imp_mun_principal, f"IMP_{ano_principal}_MUN.csv", usecols=MUN_COLS, dtypes=MUN_DTYPES)
            
            # FILTROS PRINCIPAIS PARA O DOC (AGRUPADO)
            df_exp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano, codigos_paises, agrupado, meses_para_filtrar)
            df_exp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior, codigos_paises, agrupado, meses_para_filtrar)
            
            # --- CÁLCULOS GERAIS PARA O AGRUPADO ---
            exportacao_pais_ano = df_exp_ano_mg_paises['VL_FOB'].sum()
            exportacao_pais_ano_anterior = df_exp_ano_anterior_mg_paises['VL_FOB'].sum()
            exportacao_mg_total_ano = df_exp_ano_mg['VL_FOB'].sum()
            
            df_imp_ano_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano, codigos_paises, agrupado, meses_para_filtrar)
            df_imp_ano_anterior_mg_paises = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior, codigos_paises, agrupado, meses_para_filtrar)
            importacao_pais_ano = df_imp_ano_mg_paises['VL_FOB'].sum()
            importacao_pais_ano_anterior = df_imp_ano_anterior_mg_paises['VL_FOB'].sum()
            importacao_mg_total_ano = df_imp_ano_mg['VL_FOB'].sum()
            
            balanca_ano, balanca_ano_anterior, fluxo_comercial_ano, fluxo_comercial_ano_anterior, variacao_balanca, variacao_fluxo = calcular_balanca_e_fluxo(exportacao_pais_ano, importacao_pais_ano, exportacao_pais_ano_anterior, importacao_pais_ano_anterior)

            if agrupado:
                app = DocumentoApp(logo_path=logo_path_to_use)
                paises_corretos = nomes_paises_validos 
                nome_relatorio = nome_agrupamento if (nome_agrupamento and nome_agrupamento.strip() != "") else ', '.join(paises_corretos)
                
                titulo_documento = f"Briefing - {nome_relatorio} - {ano_principal}"
                app.set_titulo(titulo_documento)
                app.nova_secao()
                
                # --- Seção 1: Balança Comercial ---
                app.adicionar_titulo("Fluxo Comercial")
                
                # Variação Fluxo
                tipo_var_fluxo = "queda" if variacao_fluxo < 0 else "acréscimo" if variacao_fluxo > 0 else "estabilidade"
                val_var_fluxo = abs(round(variacao_fluxo, 2))
                
                # Variação Balança
                tipo_var_bal = "queda" if variacao_balanca < 0 else "acréscimo" if variacao_balanca > 0 else "estabilidade"
                val_var_bal = abs(round(variacao_balanca, 2))
                saldo_str = "positiva" if balanca_ano >= 0 else "negativa"
                
                texto_balanca = (
                    f"{nome_periodo_em}, Minas Gerais e {nome_relatorio} tiveram um fluxo comercial de {formatar_valor(fluxo_comercial_ano)}, "
                    f"representando {tipo_var_fluxo} de {val_var_fluxo}% em comparação a {nome_periodo_comp}. "
                    f"A balança comercial fechou {saldo_str} para Minas Gerais em {formatar_valor(balanca_ano)}, "
                    f"apresentando uma {tipo_var_bal} de {val_var_bal}% em relação a {nome_periodo_comp}."
                )
                app.adicionar_paragrafo(texto_balanca)
                
                # --- Seção 2: Exportações DOC ---
                app.nova_secao()
                app.adicionar_titulo("Exportações")
                
                # Cálculos específicos
                ranking_mg_dest = calcular_posicao_estado_pais(df_exp_ano_estados, codigos_paises) # ranking que MG tem no BR para esse pais
                posicao_pais_para_mg_exp = 0 # Ranking que ESSE PAÍS tem para MG
                rank_df = df_exp_ano_mg.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
                # Para agrupado, é dificil dizer "O Ranking do Bloco". Vamos somar e ver onde cairia se fosse um país, ou ignorar se for bloco.
                # Se for só 1 país (agrupado=True mas len=1), calcula.
                if len(codigos_paises) == 1:
                    try:
                        posicao_pais_para_mg_exp = rank_df.index.get_loc(codigos_paises[0]) + 1
                    except:
                        posicao_pais_para_mg_exp = "-"
                else:
                     posicao_pais_para_mg_exp = "(bloco)"

                diferenca_exportacao, tipo_diferenca_exp = calcular_diferenca_percentual(exportacao_pais_ano, exportacao_pais_ano_anterior)
                participacao_pais_mg_exp = calcular_participacao(exportacao_pais_ano, exportacao_mg_total_ano)
                
                # Ranking e Part de MG no Brasil
                posicao_mg_br_exp, part_mg_br_exp = calcular_ranking_e_participacao_brasil(df_exp_ano, codigos_paises)

                # Strings de listas
                texto_produtos_exp = gerar_texto_lista_produtos(df_exp_ano_mg_paises, mapa_sh4_nomes, top_n_produtos)
                
                # Filtra municipios
                df_exp_mun_filtrado = df_exp_mun[(df_exp_mun['SG_UF_MUN'] == 'MG') & (df_exp_mun['CO_PAIS'].isin(codigos_paises)) & (df_exp_mun['CO_MES'].isin(meses_para_filtrar))]
                texto_mun_exp, count_mun_exp = gerar_texto_lista_municipios(df_exp_mun_filtrado, df_uf_mun, top_n_produtos)
                
                texto_exportacao_1 = (
                    f"{nome_relatorio} foi o {posicao_pais_para_mg_exp}º destino das exportações de Minas Gerais em {ano_principal}. "
                    f"As exportações mineiras para {nome_relatorio} somaram {formatar_valor(exportacao_pais_ano)} em {ano_principal}, "
                    f"{tipo_diferenca_exp} de {diferenca_exportacao}% em relação a {ano_comparacao}. "
                    f"A participação de {nome_relatorio} nas exportações totais de Minas Gerais em {ano_principal} foi equivalente a {participacao_pais_mg_exp}%."
                )
                app.adicionar_paragrafo(texto_exportacao_1)
                
                texto_exportacao_2 = (
                    f"Minas Gerais foi o {posicao_mg_br_exp}º principal estado exportador brasileiro para {nome_relatorio} em {ano_principal}, "
                    f"com uma participação de {part_mg_br_exp}% nas vendas do Brasil ao país."
                )
                app.adicionar_paragrafo(texto_exportacao_2)
                
                app.adicionar_paragrafo(f"Em {ano_principal}, os principais produtos exportados de Minas Gerais para {nome_relatorio} foram: {texto_produtos_exp}")
                
                app.adicionar_paragrafo(f"Dentre os {count_mun_exp} municípios de Minas Gerais que exportaram produtos para {nome_relatorio} em {ano_principal}, os principais foram: {texto_mun_exp}")

                # --- Seção 3: Importações DOC ---
                app.nova_secao()
                app.adicionar_titulo("Importações")
                
                # Cálculos imp
                diferenca_importacao, tipo_diferenca_imp = calcular_diferenca_percentual(importacao_pais_ano, importacao_pais_ano_anterior)
                participacao_pais_mg_imp = calcular_participacao(importacao_pais_ano, importacao_mg_total_ano)
                posicao_mg_br_imp, part_mg_br_imp = calcular_ranking_e_participacao_brasil(df_imp_ano, codigos_paises)
                
                rank_df_imp = df_imp_ano_mg.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
                if len(codigos_paises) == 1:
                    try:
                        posicao_pais_para_mg_imp = rank_df_imp.index.get_loc(codigos_paises[0]) + 1
                    except:
                        posicao_pais_para_mg_imp = "-"
                else:
                     posicao_pais_para_mg_imp = "(bloco)"

                texto_produtos_imp = gerar_texto_lista_produtos(df_imp_ano_mg_paises, mapa_sh4_nomes, top_n_produtos)
                
                df_imp_mun_filtrado = df_imp_mun[(df_imp_mun['SG_UF_MUN'] == 'MG') & (df_imp_mun['CO_PAIS'].isin(codigos_paises)) & (df_imp_mun['CO_MES'].isin(meses_para_filtrar))]
                texto_mun_imp, count_mun_imp = gerar_texto_lista_municipios(df_imp_mun_filtrado, df_uf_mun, top_n_produtos)

                texto_importacao_1 = (
                    f"{nome_relatorio} foi a {posicao_pais_para_mg_imp}ª origem das importações de Minas Gerais em {ano_principal}. "
                    f"As importações mineiras provenientes de {nome_relatorio} somaram {formatar_valor(importacao_pais_ano)} em {ano_principal}, "
                    f"{tipo_diferenca_imp} de {diferenca_importacao}% em relação a {ano_comparacao}. "
                    f"A participação de {nome_relatorio} nas importações totais de Minas Gerais em {ano_principal} foi equivalente a {participacao_pais_mg_imp}%."
                )
                app.adicionar_paragrafo(texto_importacao_1)
                
                texto_importacao_2 = (
                    f"Minas Gerais foi o {posicao_mg_br_imp}º principal estado importador brasileiro de {nome_relatorio} em {ano_principal}, "
                    f"com uma participação de {part_mg_br_imp}% nas compras do Brasil ao país."
                )
                app.adicionar_paragrafo(texto_importacao_2)
                
                app.adicionar_paragrafo(f"Em {ano_principal}, os principais produtos importados para Minas Gerais de {nome_relatorio} foram: {texto_produtos_imp}")
                
                app.adicionar_paragrafo(f"Dentre os {count_mun_imp} municípios de Minas Gerais que importaram produtos de {nome_relatorio} em {ano_principal}, os principais foram: {texto_mun_imp}")

                file_bytes, file_name = app.finalizar_documento() 
                st.session_state.arquivos_gerados_pais.append({"name": file_name, "data": file_bytes})
                st.success(f"Relatório '{file_name}' gerado com sucesso!")

                # --- EXIBIÇÃO APENAS NO STREAMLIT (Tabelas OTIMIZADAS) ---
                st.subheader("Visualização de Dados (Não incluído no DOCX)")
                
                def config_cols(ano_ref, ano_comp):
                    return {
                        "País": st.column_config.TextColumn("País"),
                        f"Valor {ano_ref}": st.column_config.TextColumn(f"Valor {ano_ref}"),
                        "Part. Grupo (%)": st.column_config.ProgressColumn("Part. Grupo", format="%.2f%%", min_value=0, max_value=100),
                        "Variação %": st.column_config.NumberColumn("Variação %", format="%.2f%%")
                    }

                # 1. Tabela Países (Exportação)
                st.markdown("### 🚢 Países do Bloco (Exportação)")
                grp_pais = df_exp_ano_mg_paises.groupby('CO_PAIS')['VL_FOB'].sum().reset_index().rename(columns={'VL_FOB': 'Vl_Atual'})
                grp_pais_ant = df_exp_ano_anterior_mg_paises.groupby('CO_PAIS')['VL_FOB'].sum().reset_index().rename(columns={'VL_FOB': 'Vl_Ant'})
                df_detalhe = pd.merge(grp_pais, grp_pais_ant, on='CO_PAIS', how='outer').fillna(0)
                df_detalhe['País'] = df_detalhe['CO_PAIS'].map(mapa_nomes_paises)
                df_detalhe['Variação %'] = ((df_detalhe['Vl_Atual'] - df_detalhe['Vl_Ant']) / df_detalhe['Vl_Ant'] * 100).fillna(0)
                total_grupo = df_detalhe['Vl_Atual'].sum()
                df_detalhe['Part. Grupo (%)'] = (df_detalhe['Vl_Atual'] / total_grupo * 100).fillna(0)
                df_detalhe = df_detalhe.sort_values('Vl_Atual', ascending=False)
                
                df_show = df_detalhe.copy()
                df_show[f'Valor {ano_principal}'] = df_show['Vl_Atual'].apply(formatar_valor)
                
                st.dataframe(
                    df_show[['País', f'Valor {ano_principal}', 'Part. Grupo (%)', 'Variação %']],
                    hide_index=True, use_container_width=True, column_config=config_cols(ano_principal, ano_comparacao)
                )

                # 2. PRODUTOS EXPORTAÇÃO (EXPLODIDO: LINHA POR PAÍS)
                st.caption("Top Produtos (Exportação) - Detalhado por País")
                
                # Passo A: Achar Top N Produtos (pelo Total Geral)
                top_sh4 = df_exp_ano_mg_paises.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(top_n_produtos).index.tolist()
                
                # Passo B: Filtrar dataset original apenas para esses produtos
                df_exp_top = df_exp_ano_mg_paises[df_exp_ano_mg_paises['SH4'].isin(top_sh4)].copy()
                
                # Passo C: Agrupar por Produto E País
                df_exploded_exp = df_exp_top.groupby(['SH4', 'CO_PAIS'])['VL_FOB'].sum().reset_index()
                
                # Passo D: Enriquecer
                df_exploded_exp['Produto'] = df_exploded_exp['SH4'].map(mapa_sh4_nomes)
                df_exploded_exp['País'] = df_exploded_exp['CO_PAIS'].map(mapa_nomes_paises)
                
                # Passo E: Calcular Total do Produto (para ordenar e calcular %)
                prod_totals = df_exp_top.groupby('SH4')['VL_FOB'].sum()
                df_exploded_exp['Total Produto'] = df_exploded_exp['SH4'].map(prod_totals)
                df_exploded_exp['Participação (%)'] = (df_exploded_exp['VL_FOB'] / df_exploded_exp['Total Produto']) * 100
                
                # Passo F: Ordenar (Maior Produto Primeiro -> Maior País dentro do produto)
                df_exploded_exp = df_exploded_exp.sort_values(by=['Total Produto', 'VL_FOB'], ascending=[False, False])
                
                # Passo G: Formatar
                df_exploded_exp['Valor País'] = df_exploded_exp['VL_FOB'].apply(formatar_valor)
                
                st.dataframe(
                    df_exploded_exp[['Produto', 'País', 'Valor País', 'Participação (%)']],
                    hide_index=True,
                    use_container_width=True,
                    column_config={
                        "Participação (%)": st.column_config.ProgressColumn("Share no Produto", format="%.1f%%", min_value=0, max_value=100)
                    }
                )

                st.divider()

                # 3. Tabela Países (Importação)
                st.markdown("### 📦 Países do Bloco (Importação)")
                grp_pais_imp = df_imp_ano_mg_paises.groupby('CO_PAIS')['VL_FOB'].sum().reset_index().rename(columns={'VL_FOB': 'Vl_Atual'})
                grp_pais_ant_imp = df_imp_ano_anterior_mg_paises.groupby('CO_PAIS')['VL_FOB'].sum().reset_index().rename(columns={'VL_FOB': 'Vl_Ant'})
                df_detalhe_imp = pd.merge(grp_pais_imp, grp_pais_ant_imp, on='CO_PAIS', how='outer').fillna(0)
                df_detalhe_imp['País'] = df_detalhe_imp['CO_PAIS'].map(mapa_nomes_paises)
                df_detalhe_imp['Variação %'] = ((df_detalhe_imp['Vl_Atual'] - df_detalhe_imp['Vl_Ant']) / df_detalhe_imp['Vl_Ant'] * 100).fillna(0)
                total_grupo_imp = df_detalhe_imp['Vl_Atual'].sum()
                df_detalhe_imp['Part. Grupo (%)'] = (df_detalhe_imp['Vl_Atual'] / total_grupo_imp * 100).fillna(0)
                df_detalhe_imp = df_detalhe_imp.sort_values('Vl_Atual', ascending=False)
                
                df_show_imp = df_detalhe_imp.copy()
                df_show_imp[f'Valor {ano_principal}'] = df_show_imp['Vl_Atual'].apply(formatar_valor)
                
                st.dataframe(
                    df_show_imp[['País', f'Valor {ano_principal}', 'Part. Grupo (%)', 'Variação %']],
                    hide_index=True, use_container_width=True, column_config=config_cols(ano_principal, ano_comparacao)
                )

                # 4. PRODUTOS IMPORTAÇÃO (EXPLODIDO)
                st.caption("Top Produtos (Importação) - Detalhado por País")
                
                top_sh4_imp = df_imp_ano_mg_paises.groupby('SH4')['VL_FOB'].sum().sort_values(ascending=False).head(top_n_produtos).index.tolist()
                df_imp_top = df_imp_ano_mg_paises[df_imp_ano_mg_paises['SH4'].isin(top_sh4_imp)].copy()
                df_exploded_imp = df_imp_top.groupby(['SH4', 'CO_PAIS'])['VL_FOB'].sum().reset_index()
                
                df_exploded_imp['Produto'] = df_exploded_imp['SH4'].map(mapa_sh4_nomes)
                df_exploded_imp['País'] = df_exploded_imp['CO_PAIS'].map(mapa_nomes_paises)
                
                prod_totals_imp = df_imp_top.groupby('SH4')['VL_FOB'].sum()
                df_exploded_imp['Total Produto'] = df_exploded_imp['SH4'].map(prod_totals_imp)
                df_exploded_imp['Participação (%)'] = (df_exploded_imp['VL_FOB'] / df_exploded_imp['Total Produto']) * 100
                
                df_exploded_imp = df_exploded_imp.sort_values(by=['Total Produto', 'VL_FOB'], ascending=[False, False])
                df_exploded_imp['Valor País'] = df_exploded_imp['VL_FOB'].apply(formatar_valor)
                
                st.dataframe(
                    df_exploded_imp[['Produto', 'País', 'Valor País', 'Participação (%)']],
                    hide_index=True,
                    use_container_width=True,
                    column_config={
                        "Participação (%)": st.column_config.ProgressColumn("Share no Produto", format="%.1f%%", min_value=0, max_value=100)
                    }
                )

            else:
                # --- LÓGICA PARA SEPARADOS ---
                paises_corretos = nomes_paises_validos
                
                df_exp_ano_loop = df_exp_ano.copy()
                df_exp_ano_anterior_loop = df_exp_ano_anterior.copy()
                df_imp_ano_loop = df_imp_ano.copy()
                df_imp_ano_anterior_loop = df_imp_ano_anterior.copy()

                for pais in paises_corretos:
                    st.subheader(f"Processando: {pais}") 
                    app = DocumentoApp(logo_path=logo_path_to_use)
                    
                    codigos_paises_loop = [obter_codigo_pais(pais, mapa_paises_reverso)]

                    # Filtra dados para este país
                    df_exp_atual_loop = filtrar_dados_por_mg_e_pais(df_exp_ano_loop, codigos_paises_loop, False, meses_para_filtrar)
                    df_exp_ant_loop = filtrar_dados_por_mg_e_pais(df_exp_ano_anterior_loop, codigos_paises_loop, False, meses_para_filtrar)
                    df_imp_atual_loop = filtrar_dados_por_mg_e_pais(df_imp_ano_loop, codigos_paises_loop, False, meses_para_filtrar)
                    df_imp_ant_loop = filtrar_dados_por_mg_e_pais(df_imp_ano_anterior_loop, codigos_paises_loop, False, meses_para_filtrar)
                    
                    # Valores e Balança
                    v_exp_atual = df_exp_atual_loop['VL_FOB'].sum()
                    v_exp_ant = df_exp_ant_loop['VL_FOB'].sum()
                    v_imp_atual = df_imp_atual_loop['VL_FOB'].sum()
                    v_imp_ant = df_imp_ant_loop['VL_FOB'].sum()
                    
                    balanca_loop, balanca_ant_loop, fluxo_loop, fluxo_ant_loop, var_bal, var_fluxo = calcular_balanca_e_fluxo(v_exp_atual, v_imp_atual, v_exp_ant, v_imp_ant)

                    titulo_documento = f"Briefing - {pais} - {ano_principal}"
                    app.set_titulo(titulo_documento)
                    app.nova_secao()

                    # Texto Balança
                    app.adicionar_titulo("Fluxo Comercial")
                    tipo_var_fluxo = "queda" if var_fluxo < 0 else "acréscimo" if var_fluxo > 0 else "estabilidade"
                    val_var_fluxo = abs(round(var_fluxo, 2))
                    tipo_var_bal = "queda" if var_bal < 0 else "acréscimo" if var_bal > 0 else "estabilidade"
                    val_var_bal = abs(round(var_bal, 2))
                    saldo_str = "positiva" if balanca_loop >= 0 else "negativa"
                    
                    texto_balanca = (
                        f"{nome_periodo_em}, Minas Gerais e {pais} tiveram um fluxo comercial de {formatar_valor(fluxo_loop)}, "
                        f"representando {tipo_var_fluxo} de {val_var_fluxo}% em comparação a {nome_periodo_comp}. "
                        f"A balança comercial fechou {saldo_str} para Minas Gerais em {formatar_valor(balanca_loop)}, "
                        f"apresentando uma {tipo_var_bal} de {val_var_bal}% em relação a {nome_periodo_comp}."
                    )
                    app.adicionar_paragrafo(texto_balanca)
                    
                    # Texto Exportações
                    app.nova_secao()
                    app.adicionar_titulo("Exportações")
                    
                    # Rankings
                    rank_df = df_exp_ano_mg.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
                    try:
                        pos_pais_mg = rank_df.index.get_loc(codigos_paises_loop[0]) + 1
                    except:
                        pos_pais_mg = "-"
                    
                    dif_exp_val, tipo_dif_exp = calcular_diferenca_percentual(v_exp_atual, v_exp_ant)
                    part_exp = calcular_participacao(v_exp_atual, exportacao_mg_total_ano)
                    pos_mg_br_exp, part_mg_br_exp = calcular_ranking_e_participacao_brasil(df_exp_ano, codigos_paises_loop)
                    
                    texto_prods_exp = gerar_texto_lista_produtos(df_exp_atual_loop, mapa_sh4_nomes, top_n_produtos)
                    
                    df_exp_mun_loop = df_exp_mun[(df_exp_mun['SG_UF_MUN'] == 'MG') & (df_exp_mun['CO_PAIS'].isin(codigos_paises_loop)) & (df_exp_mun['CO_MES'].isin(meses_para_filtrar))]
                    texto_mun_exp, count_mun_exp = gerar_texto_lista_municipios(df_exp_mun_loop, df_uf_mun, top_n_produtos)

                    app.adicionar_paragrafo(
                        f"{pais} foi o {pos_pais_mg}º destino das exportações de Minas Gerais em {ano_principal}. "
                        f"As exportações mineiras para {pais} somaram {formatar_valor(v_exp_atual)} em {ano_principal}, "
                        f"{tipo_dif_exp} de {dif_exp_val}% em relação a {ano_comparacao}. "
                        f"A participação de {pais} nas exportações totais de Minas Gerais em {ano_principal} foi equivalente a {part_exp}%."
                    )
                    app.adicionar_paragrafo(
                        f"Minas Gerais foi o {pos_mg_br_exp}º principal estado exportador brasileiro para {pais} em {ano_principal}, "
                        f"com uma participação de {part_mg_br_exp}% nas vendas do Brasil ao país."
                    )
                    app.adicionar_paragrafo(f"Em {ano_principal}, os principais produtos exportados de Minas Gerais para {pais} foram: {texto_prods_exp}")
                    app.adicionar_paragrafo(f"Dentre os {count_mun_exp} municípios de Minas Gerais que exportaram produtos para {pais} em {ano_principal}, os principais foram: {texto_mun_exp}")

                    # Texto Importações
                    app.nova_secao()
                    app.adicionar_titulo("Importações")
                    
                    rank_df_imp = df_imp_ano_mg.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
                    try:
                        pos_pais_mg_imp = rank_df_imp.index.get_loc(codigos_paises_loop[0]) + 1
                    except:
                        pos_pais_mg_imp = "-"
                    
                    dif_imp_val, tipo_dif_imp = calcular_diferenca_percentual(v_imp_atual, v_imp_ant)
                    part_imp = calcular_participacao(v_imp_atual, importacao_mg_total_ano)
                    pos_mg_br_imp, part_mg_br_imp = calcular_ranking_e_participacao_brasil(df_imp_ano, codigos_paises_loop)
                    
                    texto_prods_imp = gerar_texto_lista_produtos(df_imp_atual_loop, mapa_sh4_nomes, top_n_produtos)
                    
                    df_imp_mun_loop = df_imp_mun[(df_imp_mun['SG_UF_MUN'] == 'MG') & (df_imp_mun['CO_PAIS'].isin(codigos_paises_loop)) & (df_imp_mun['CO_MES'].isin(meses_para_filtrar))]
                    texto_mun_imp, count_mun_imp = gerar_texto_lista_municipios(df_imp_mun_loop, df_uf_mun, top_n_produtos)

                    app.adicionar_paragrafo(
                        f"{pais} foi a {pos_pais_mg_imp}ª origem das importações de Minas Gerais em {ano_principal}. "
                        f"As importações mineiras provenientes de {pais} somaram {formatar_valor(v_imp_atual)} em {ano_principal}, "
                        f"{tipo_dif_imp} de {dif_imp_val}% em relação a {ano_comparacao}. "
                        f"A participação de {pais} nas importações totais de Minas Gerais em {ano_principal} foi equivalente a {part_imp}%."
                    )
                    app.adicionar_paragrafo(
                        f"Minas Gerais foi o {pos_mg_br_imp}º principal estado importador brasileiro de {pais} em {ano_principal}, "
                        f"com uma participação de {part_mg_br_imp}% nas compras do Brasil ao país."
                    )
                    app.adicionar_paragrafo(f"Em {ano_principal}, os principais produtos importados para Minas Gerais de {pais} foram: {texto_prods_imp}")
                    app.adicionar_paragrafo(f"Dentre os {count_mun_imp} municípios de Minas Gerais que importaram produtos de {pais} em {ano_principal}, os principais foram: {texto_mun_imp}")

                    file_bytes, file_name = app.finalizar_documento()
                    st.session_state.arquivos_gerados_pais.append({"name": file_name, "data": file_bytes})
            
            # Limpeza de memória
            del df_exp_ano, df_exp_ano_anterior, df_imp_ano, df_imp_ano_anterior
            if 'df_exp_ano_loop' in locals():
                del df_exp_ano_loop, df_exp_ano_anterior_loop, df_imp_ano_loop, df_imp_ano_anterior_loop
                
        except Exception as e:
            st.error(f"Ocorreu um erro inesperado durante a geração:")
            st.exception(e)

# --- Bloco de exibição de Download ---
if st.session_state.arquivos_gerados_pais:
    st.header("4. Relatórios Gerados")
    st.info("Clique para baixar os relatórios. Eles permanecerão aqui até que você gere um novo relatório.")
    
    if len(st.session_state.arquivos_gerados_pais) > 1:
        st.subheader("Pacote de Relatórios (ZIP)")
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for arquivo in st.session_state.arquivos_gerados_pais:
                zip_file.writestr(arquivo["name"], arquivo["data"])
        
        zip_bytes = zip_buffer.getvalue()
        
        st.download_button(
            label=f"Baixar todos os {len(st.session_state.arquivos_gerados_pais)} relatórios (.zip)",
            data=zip_bytes,
            file_name=f"Briefings_Países_{ano_principal}.zip",
            mime="application/zip",
            key="download_zip_pais"
        )
        
    elif len(st.session_state.arquivos_gerados_pais) == 1:
        st.subheader("Relatório Gerado")
        arquivo = st.session_state.arquivos_gerados_pais[0] 
        st.download_button(
            label=f"Baixar Relatório ({arquivo['name']})",
            data=arquivo["data"], 
            file_name=arquivo["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{arquivo['name']}"
        )

# --- Bloco de Rodapé ---
st.divider() 

col1, col2 = st.columns([0.3, 0.7], vertical_alignment="center") 

with col1:
    logo_footer_path = "AEST Sede.png"
    if os.path.exists(logo_footer_path):
        st.image(logo_footer_path, width=150)
    else:
        st.caption("Aest - Dados")

with col2:
    st.caption("Desenvolvido por Aest - Dados e Subsecretaria de Promoção de Investimentos e Cadeias Produtivas")