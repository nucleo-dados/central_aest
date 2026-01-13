import streamlit as st
import pandas as pd
import requests
from io import StringIO
from urllib3.exceptions import InsecureRequestWarning
import os
from datetime import datetime
import io
import re
import zipfile
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- IMPORTAÇÃO E PROTEÇÃO DA PÁGINA ---
try:
    # --- ALTERAÇÃO: Importa do auth.py ---
    from auth import page_protector 
except ImportError:
    st.error("Erro ao importar autenticação. Execute a partir do Home.py")
    st.stop()

# Protege a página, oculta nav padrão e desenha a sidebar
page_protector(page_name="Análise por País")
# --- FIM DA PROTEÇÃO ---

# --- CONFIGURAÇÕES GLOBAIS ---
requests.packages.urllib3.disable_warnings(category=InsecureRequestWarning)

MESES_MAPA = {
    "Janeiro": 1, "Fevereiro": 2, "Março": 3, "Abril": 4, "Maio": 5, "Junho": 6,
    "Julho": 7, "Agosto": 8, "Setembro": 9, "Outubro": 10, "Novembro": 11, "Dezembro": 12
}
LISTA_MESES = list(MESES_MAPA.keys())
meses_pt = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
    7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
}

# Colunas necessárias
NCM_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_NCM', 'CO_NCM']
NCM_DTYPES = {'CO_NCM': str, 'CO_SH4': str}
# --- NOVO: Colunas e dtypes para arquivos MUN (Embora não possam ser usados para produto) ---
MUN_COLS = ['VL_FOB', 'CO_PAIS', 'CO_MES', 'SG_UF_MUN', 'CO_MUN']
MUN_DTYPES = {'CO_MUN': str}

# --- FUNÇÕES DE LÓGICA (Helpers) ---

@st.cache_data(ttl=3600)
def ler_dados_csv_online(url, usecols=None, dtypes=None):
    """Lê dados CSV da URL com retentativas."""
    retries = 3
    for attempt in range(retries):
        try:
            resposta = requests.get(url, verify=False, timeout=(10, 1200)) 
            resposta.raise_for_status()
            df = pd.read_csv(StringIO(resposta.content.decode('latin-1')), encoding='latin-1',
                             sep=';', dtype=dtypes, usecols=usecols)
            return df
        except requests.exceptions.RequestException as e:
            st.error(f"Erro ao acessar o CSV (tentativa {attempt + 1}/{retries}): {e}")
            if "Read timed out" in str(e) and attempt < retries - 1:
                st.warning("Download demorou muito. Tentando novamente...")
                continue
            else:
                st.error(f"Falha ao baixar após {retries} tentativas.")
                return None
        except Exception as e:
            st.error(f"Erro inesperado ao baixar ou processar o CSV: {e}")
            return None
    return None

@st.cache_data(ttl=3600)
def carregar_dataframe(url, nome_arquivo, usecols=None, dtypes=None, mostrar_progresso=True):
    """Carrega o DataFrame da URL (usa cache) com colunas e dtypes."""
    progress_bar = None
    # --- ALTERADO: Oculta a barra de progresso para os arquivos de UFs e Municípios ---
    if mostrar_progresso and nome_arquivo not in ["EXP_UF_TODOS", "IMP_UF_TODOS", "EXP_MUN", "IMP_MUN", "UF_MUN.csv"]: 
        progress_bar = st.progress(0, text=f"Carregando {nome_arquivo}...")
    
    df = ler_dados_csv_online(url, usecols=usecols, dtypes=dtypes)
    
    if mostrar_progresso and progress_bar: 
        if df is not None:
            progress_bar.progress(100, text=f"{nome_arquivo} carregado com sucesso.")
        else:
            progress_bar.empty()
    return df

@st.cache_data
def obter_dados_paises():
    """Carrega a tabela de países (ID e Nome) e armazena em cache."""
    url_pais = "https://balanca.economia.gov.br/balanca/bd/tabelas/PAIS.csv"
    df_pais = carregar_dataframe(url_pais, "PAIS.csv", usecols=['NO_PAIS', 'CO_PAIS'], mostrar_progresso=False) 
    if df_pais is not None and not df_pais.empty:
        mapa_codigo_nome = pd.Series(df_pais.NO_PAIS.values, index=df_pais.CO_PAIS).to_dict()
        lista_nomes = sorted(df_pais[df_pais['NO_PAIS'] != 'Brasil']['NO_PAIS'].unique().tolist())
        mapa_nome_codigo = pd.Series(df_pais.CO_PAIS.values, index=df_pais.NO_PAIS).to_dict()
        return mapa_codigo_nome, lista_nomes, mapa_nome_codigo
    return {}, [], {}

@st.cache_data
def obter_dados_produtos_ncm():
    """Carrega a tabela NCM completa (SH2, SH4 e SH6) e armazena em cache."""
    url_ncm = "https://balanca.economia.gov.br/balanca/bd/tabelas/NCM_SH.csv"
    usecols_ncm = ['CO_SH2', 'NO_SH2_POR', 'CO_SH4', 'NO_SH4_POR', 'CO_SH6', 'NO_SH6_POR']
    df_ncm = carregar_dataframe(url_ncm, "NCM_SH.csv", usecols=usecols_ncm, mostrar_progresso=False)
    if df_ncm is not None:
        # Criar mapas de nomes de produtos para reuso
        df_ncm['CO_SH2_STR'] = df_ncm['CO_SH2'].astype(str).str.zfill(2)
        df_ncm['CO_SH4_STR'] = df_ncm['CO_SH4'].astype(str).str.zfill(4)
        df_ncm['CO_SH6_STR'] = df_ncm['CO_SH6'].astype(str).str.zfill(6)
        
        mapa_sh2 = df_ncm.drop_duplicates('CO_SH2_STR').set_index('CO_SH2_STR')['NO_SH2_POR']
        mapa_sh4 = df_ncm.drop_duplicates('CO_SH4_STR').set_index('CO_SH4_STR')['NO_SH4_POR']
        mapa_sh6 = df_ncm.drop_duplicates('CO_SH6_STR').set_index('CO_SH6_STR')['NO_SH6_POR']
        
        return df_ncm, mapa_sh2.to_dict(), mapa_sh4.to_dict(), mapa_sh6.to_dict()
    return None, {}, {}, {}

def obter_lista_de_produtos_sh2():
    """Retorna uma lista de capítulos (SH2)."""
    df_ncm, _, _, _ = obter_dados_produtos_ncm()
    if df_ncm is not None:
        df_sh2 = df_ncm.drop_duplicates(subset=['CO_SH2']).dropna()
        df_sh2['Display'] = df_sh2['CO_SH2'].astype(str).str.zfill(2) + " - " + df_sh2['NO_SH2_POR']
        lista_produtos = df_sh2['Display'].unique().tolist()
        lista_produtos.sort()
        return lista_produtos
    return ["Erro ao carregar lista de capítulos"]

def obter_lista_de_produtos_sh4():
    """Retorna uma lista de produtos (SH4)."""
    df_ncm, _, _, _ = obter_dados_produtos_ncm()
    if df_ncm is None:
        return ["Erro ao carregar lista de produtos"]

    df_sh4 = df_ncm.drop_duplicates(subset=['CO_SH4']).dropna(subset=['CO_SH4', 'NO_SH4_POR'])
    df_sh4['Display'] = df_sh4['CO_SH4'].astype(str).str.zfill(4) + " - " + df_sh4['NO_SH4_POR']
    lista_produtos = df_sh4['Display'].unique().tolist()
    lista_produtos.sort()
    return lista_produtos

@st.cache_data
def obter_lista_de_produtos_sh6():
    """Retorna uma lista de produtos (SH6)."""
    df_ncm, _, _, _ = obter_dados_produtos_ncm()
    if df_ncm is None:
        return ["Erro ao carregar lista de SH6"]
    
    df_sh6 = df_ncm.drop_duplicates(subset=['CO_SH6']).dropna(subset=['CO_SH6', 'NO_SH6_POR'])
    df_sh6['Display'] = df_sh6['CO_SH6'].astype(str).str.zfill(6) + " - " + df_sh6['NO_SH6_POR']
    lista_produtos = df_sh6['Display'].unique().tolist()
    lista_produtos.sort()
    return lista_produtos

def get_sh2(co_ncm):
    """Extrai SH2 de um CO_NCM."""
    co_ncm_str = str(co_ncm).strip()
    if pd.isna(co_ncm_str) or co_ncm_str == "":
        return None
    co_ncm_str = co_ncm_str.zfill(8)
    return co_ncm_str[:2]

def get_sh4(co_ncm):
    """Extrai SH4 de um CO_NCM."""
    co_ncm_str = str(co_ncm).strip()
    if pd.isna(co_ncm_str) or co_ncm_str == "":
        return None
    co_ncm_str = co_ncm_str.zfill(8)
    return co_ncm_str[:4]

def get_sh6(co_ncm):
    """Extrai SH6 de um CO_NCM."""
    co_ncm_str = str(co_ncm).strip()
    if pd.isna(co_ncm_str) or co_ncm_str == "":
        return None
    co_ncm_str = co_ncm_str.zfill(8)
    return co_ncm_str[:6] 

def formatar_valor(valor):
    prefixo = ""
    if valor < 0:
        prefixo = "-"
        valor = abs(valor)
    if valor >= 1_000_000_000:
        valor_formatado_str = f"{(valor / 1_000_000_000):.2f}".replace('.',',')
        unidade = "bilhão" if (valor / 1_000_000_000) < 2 else "bilhões"
        return f"{prefixo}US$ {valor_formatado_str} {unidade}"
    if valor >= 1_000_000:
        valor_formatado_str = f"{(valor / 1_000_000):.2f}".replace('.',',')
        unidade = "milhão" if (valor / 1_000_000) < 2 else "milhões"
        return f"{prefixo}US$ {valor_formatado_str} {unidade}"
    if valor >= 1_000:
        valor_formatado_str = f"{(valor / 1_000):.2f}".replace('.',',')
        return f"{prefixo}US$ {valor_formatado_str} mil"
    valor_formatado_str = f"{valor:.2f}".replace('.',',')
    return f"{prefixo}US$ {valor_formatado_str}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def calcular_diferenca_percentual(valor_atual, valor_anterior):
    """Calcula a diferença percentual entre dois valores."""
    if valor_anterior == 0:
        return 0.0, "acréscimo" if valor_atual > 0 else "redução" if valor_atual < 0 else "estabilidade"
    diferenca = round(((valor_atual - valor_anterior) / valor_anterior) * 100, 2)
    if diferenca > 0:
        tipo_diferenca = "um acréscimo"
    elif diferenca < 0:
        tipo_diferenca = "uma redução"
    else:
        tipo_diferenca = "uma estabilidade"
    diferenca = abs(diferenca)
    return diferenca, tipo_diferenca

class DocumentoApp:
    def __init__(self, logo_path):
        self.doc = Document()
        self.secao_atual = 0
        self.subsecao_atual = 0
        self.titulo_doc = ""
        self.logo_path = logo_path
        self.diretorio_base = "/tmp/" 
    def set_titulo(self, titulo):
        self.titulo_doc = sanitize_filename(titulo)
        self.criar_cabecalho()
        p = self.doc.add_paragraph()
        run = p.add_run(self.titulo_doc)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def adicionar_conteudo_formatado(self, texto):
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    def adicionar_paragrafo(self, texto): 
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        run = p.add_run(texto)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    def adicionar_titulo(self, texto):
        p = self.doc.add_paragraph()
        if self.subsecao_atual == 0:
            run = p.add_run(f"{self.secao_atual}. {texto}")
        else:
            run = p.add_run(f"{self.secao_atual}.{self.subsecao_atual}. {texto}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def nova_secao(self):
        self.secao_atual += 1
        self.subsecao_atual = 0
    def criar_cabecalho(self):
        section = self.doc.sections[0]
        section.top_margin = Cm(1.27)
        header = section.header
        largura_total_cm = 16.0
        table = header.add_table(rows=1, cols=2, width=Cm(largura_total_cm))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.columns[0].width = Cm(4.0)
        table.columns[1].width = Cm(12.0)
        cell_imagem = table.cell(0, 0)
        paragraph_imagem = cell_imagem.paragraphs[0]
        paragraph_imagem.paragraph_format.space_before = Pt(0)
        paragraph_imagem.paragraph_format.space_after = Pt(0)
        run_imagem = paragraph_imagem.add_run()
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                run_imagem.add_picture(self.logo_path,
                                       width=Cm(3.5), 
                                       height=Cm(3.42))
            except Exception as e:
                paragraph_imagem.add_run("[Logo não encontrado]")
        else:
            paragraph_imagem.add_run("[Logo não encontrado]")
        paragraph_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_texto = table.cell(0, 1)
        textos = [
            "GOVERNO DO ESTADO DE MINAS GERAIS",
            "SECRETARIA DE ESTADO DE DESENVOLVIMENTO ECONÔMICO",
            "Subsecretaria de Promoção de Investimentos e Cadeias Produtivas",
            "Superintendência de Atração de Investimentos e Estímulo à Exportação"
        ]
        def formatar_paragrafo_cabecalho(p):
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p = cell_texto.paragraphs[0]
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[0])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True 
        p = cell_texto.add_paragraph()
        formatar_paragrafo_cabecalho(p)
        run = p.add_run(textos[1])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        for texto in textos[2:]: 
            p = cell_texto.add_paragraph()
            formatar_paragrafo_cabecalho(p)
            run = p.add_run(texto)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = False 
    def finalizar_documento(self):
        diretorio_real = self.diretorio_base
        try:
            os.makedirs(diretorio_real, exist_ok=True)
        except Exception:
            diretorio_real = "/tmp/"
            os.makedirs(diretorio_real, exist_ok=True)
        nome_arquivo = f"{self.titulo_doc}.docx"
        nome_arquivo_sanitizado = sanitize_filename(nome_arquivo)
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        file_bytes = file_stream.getvalue()
        st.success(f"Documento '{nome_arquivo_sanitizado}' gerado com sucesso!")
        return file_bytes, nome_arquivo_sanitizado

# --- CONFIGURAÇÃO DA PÁGINA ---

st.header("1. Configurações da Análise de Produto (NCM)")

# --- Callback para limpar o state ---
def clear_download_state_prod():
    """Limpa os relatórios gerados da sessão."""
    if 'arquivos_gerados_produto' in st.session_state:
        st.session_state.arquivos_gerados_produto = []

# Carrega dados de Países e Produtos
lista_de_produtos_sh2 = obter_lista_de_produtos_sh2()
lista_de_produtos_sh4 = obter_lista_de_produtos_sh4()
lista_de_produtos_sh6 = obter_lista_de_produtos_sh6()
mapa_nomes_paises, lista_paises_nomes, mapa_paises_reverso = obter_dados_paises()
df_ncm_completo, mapa_sh2_nomes, mapa_sh4_nomes, mapa_sh6_nomes = obter_dados_produtos_ncm() 
ano_atual = datetime.now().year

col1, col2 = st.columns(2)
with col1:
    ano_principal = st.number_input(
        "Ano de Referência:", min_value=1998, max_value=ano_atual, value=ano_atual,
        help="O ano principal que você quer analisar.",
        on_change=clear_download_state_prod 
    )
    ano_comparacao = st.number_input(
        "Ano de Comparação:", min_value=1998, max_value=ano_atual, value=ano_atual - 1,
        help="O ano contra o qual você quer comparar.",
        on_change=clear_download_state_prod 
    )
    meses_selecionados = st.multiselect(
        "Meses de Análise (opcional):",
        options=LISTA_MESES,
        help="Selecione os meses. Se deixar em branco, o ano inteiro será analisado.",
        on_change=clear_download_state_prod 
    )
    
    top_n_paises = st.number_input(
        "Nº de Países no Ranking:",
        min_value=1,
        max_value=100,
        value=10,
        help="Quantos países devem ser exibidos nas tabelas de ranking (Top 10, Top 20, etc.).",
        on_change=clear_download_state_prod
    )

with col2:
    paises_selecionados_nomes = st.multiselect(
        "Filtrar por País (opcional):",
        options=lista_paises_nomes,
        help="Filtre a análise para países específicos (destino na EXP, origem na IMP).",
        on_change=clear_download_state_prod 
    )

    sh2_selecionados_nomes = st.multiselect(
        "1. Selecione Capítulos (SH2) (opcional):",
        options=lista_de_produtos_sh2,
        help="Selecione um ou mais capítulos (2 dígitos).",
        on_change=clear_download_state_prod 
    )
    
    sh4_selecionados_nomes = st.multiselect(
        "2. Selecione Produtos (SH4) (opcional):",
        options=lista_de_produtos_sh4,
        default=[],
        help="Selecione um ou mais produtos (4 dígitos).",
        on_change=clear_download_state_prod 
    )
    
    sh6_selecionados_nomes = st.multiselect(
        "3. Selecione Subposições (SH6) (opcional):",
        options=lista_de_produtos_sh6,
        default=[],
        help="Selecione uma ou mais subposições (6 dígitos).",
        on_change=clear_download_state_prod
    )

# --- Lógica de Agrupamento ---
agrupado = True
nome_agrupamento = None 

total_selecionado = len(sh2_selecionados_nomes) + len(sh4_selecionados_nomes) + len(sh6_selecionados_nomes)
produtos_para_agrupar_nomes = sh2_selecionados_nomes + sh4_selecionados_nomes + sh6_selecionados_nomes

if total_selecionado > 1:
    st.header("2. Opções de Agrupamento")
    agrupamento_input = st.radio(
        f"Deseja que os dados dos {total_selecionado} produtos/grupos sejam agrupados?",
        ("agrupados", "separados"),
        index=0,
        horizontal=True,
        on_change=clear_download_state_prod 
    )
    agrupado = (agrupamento_input == "agrupados")
    
    if agrupado:
        st.info(
            "💡 **Como funciona o agrupamento:**\n"
            "* **Agrupados:** Gerará um **único relatório** consolidado. As tabelas de ranking mostrarão a **soma** de todos os produtos (SH2, SH4, SH6) selecionados. O detalhamento por produto aparecerá no expansor.\n"
            "* **Separados:** Gerará um **relatório individual** para cada item que você selecionou (ex: um relatório para o SH2, um para o SH4, etc.)."
        )
        
        quer_nome_agrupamento = st.checkbox(
            "Deseja dar um nome para este agrupamento de produtos?", 
            key="prod_nome_grupo",
            on_change=clear_download_state_prod
        )
        if quer_nome_agrupamento:
            nome_agrupamento = st.text_input(
                "Digite o nome do agrupamento:", 
                key="prod_nome_input",
                on_change=clear_download_state_prod
            )
    
    st.header("3. Gerar Análise")
else:
    agrupado = False
    st.header("2. Gerar Análise")

# --- Inicialização do Session State ---
if 'arquivos_gerados_produto' not in st.session_state:
    st.session_state.arquivos_gerados_produto = []


if st.button("Iniciar Análise por Produto"):
    
    st.session_state.arquivos_gerados_produto = []
    logo_path_to_use = "LogoMinasGerais.png" 
    
    with st.spinner(f"Processando dados de produto..."):
        try:
            codigos_sh2_selecionados = [s.split(" - ")[0] for s in sh2_selecionados_nomes]
            codigos_sh4_selecionados = [s.split(" - ")[0] for s in sh4_selecionados_nomes]
            codigos_sh6_selecionados = [s.split(" - ")[0] for s in sh6_selecionados_nomes]
            
            if not codigos_sh2_selecionados and not codigos_sh4_selecionados and not codigos_sh6_selecionados:
                st.error("Nenhum produto (SH2, SH4 ou SH6) selecionado.")
                st.stop()
            
            codigos_paises_selecionados = [mapa_paises_reverso[nome] for nome in paises_selecionados_nomes]

            # --- NOVO: URLs de Município ---
            url_exp_ano_principal_mun = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/EXP_{ano_principal}_MUN.csv"
            url_imp_ano_principal_mun = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/mun/IMP_{ano_principal}_MUN.csv"
            url_uf_mun = "https://balanca.economia.gov.br/balanca/bd/tabelas/UF_MUN.csv"
            
            # --- ATENÇÃO: Carregando dados de TODAS AS UFs para o ranking nacional ---
            # (mostrar_progresso=False para não poluir a UI)
            url_exp_ano_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_principal}.csv"
            url_exp_ano_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/EXP_{ano_comparacao}.csv"
            url_imp_ano_principal = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_principal}.csv"
            url_imp_ano_comparacao = f"https://balanca.economia.gov.br/balanca/bd/comexstat-bd/ncm/IMP_{ano_comparacao}.csv"

            df_exp_princ_ufs = carregar_dataframe(url_exp_ano_principal, "EXP_UF_TODOS", usecols=NCM_COLS, dtypes=NCM_DTYPES, mostrar_progresso=False)
            df_exp_comp_ufs = carregar_dataframe(url_exp_ano_comparacao, "EXP_UF_TODOS", usecols=NCM_COLS, dtypes=NCM_DTYPES, mostrar_progresso=False)
            df_imp_princ_ufs = carregar_dataframe(url_imp_ano_principal, "IMP_UF_TODOS", usecols=NCM_COLS, dtypes=NCM_DTYPES, mostrar_progresso=False)
            df_imp_comp_ufs = carregar_dataframe(url_imp_ano_comparacao, "IMP_UF_TODOS", usecols=NCM_COLS, dtypes=NCM_DTYPES, mostrar_progresso=False)

            # --- NOVO: Carregamento de dados MUN ---
            # (Aviso: Estes arquivos não contêm CO_NCM, então o ranking municipal é impossível)
            # df_exp_mun_princ = carregar_dataframe(url_exp_ano_principal_mun, "EXP_MUN", usecols=MUN_COLS, dtypes=MUN_DTYPES, mostrar_progresso=False)
            # df_imp_mun_princ = carregar_dataframe(url_imp_ano_principal_mun, "IMP_MUN", usecols=MUN_COLS, dtypes=MUN_DTYPES, mostrar_progresso=False)
            # df_uf_mun = carregar_dataframe(url_uf_mun, "UF_MUN.csv", usecols=['CO_MUN_GEO', 'NO_MUN_MIN'], mostrar_progresso=False)
            
            # Verificação de falha
            if df_exp_princ_ufs is None or df_imp_princ_ufs is None or df_exp_comp_ufs is None or df_imp_comp_ufs is None:
                st.error("Falha ao carregar arquivos de dados NCM (Nacional). Tente novamente.")
                st.stop()
            # if df_exp_mun_princ is None or df_imp_mun_princ is None or df_uf_mun is None:
            #     st.error("Falha ao carregar arquivos de dados Municipais. Tente novamente.")
            #     st.stop()
            st.warning("AVISO: Os arquivos públicos da Comex Stat não permitem cruzar dados de Produto (NCM) com Município. O ranking municipal não será gerado.")
            
            
            # --- Filtro de Meses ---
            if meses_selecionados:
                meses_para_filtrar = [MESES_MAPA[m] for m in meses_selecionados]
                nome_periodo = f"o período de {', '.join(meses_selecionados)} de {ano_principal}"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            else:
                ultimo_mes_disponivel = df_exp_princ_ufs['CO_MES'].max()
                meses_para_filtrar = list(range(1, ultimo_mes_disponivel + 1))
                nome_periodo = f"o ano de {ano_principal} (até {meses_pt.get(ultimo_mes_disponivel, ultimo_mes_disponivel)})"
                nome_periodo_comp = f"o mesmo período de {ano_comparacao}"
            
            # --- Adiciona colunas SH (Nível UF) ---
            df_exp_princ_ufs['SH2'] = df_exp_princ_ufs['CO_NCM'].apply(get_sh2)
            df_exp_comp_ufs['SH2'] = df_exp_comp_ufs['CO_NCM'].apply(get_sh2)
            df_imp_princ_ufs['SH2'] = df_imp_princ_ufs['CO_NCM'].apply(get_sh2)
            df_imp_comp_ufs['SH2'] = df_imp_comp_ufs['CO_NCM'].apply(get_sh2)
            
            df_exp_princ_ufs['SH4'] = df_exp_princ_ufs['CO_NCM'].apply(get_sh4)
            df_exp_comp_ufs['SH4'] = df_exp_comp_ufs['CO_NCM'].apply(get_sh4)
            df_imp_princ_ufs['SH4'] = df_imp_princ_ufs['CO_NCM'].apply(get_sh4)
            df_imp_comp_ufs['SH4'] = df_imp_comp_ufs['CO_NCM'].apply(get_sh4)
            
            df_exp_princ_ufs['SH6'] = df_exp_princ_ufs['CO_NCM'].apply(get_sh6)
            df_exp_comp_ufs['SH6'] = df_exp_comp_ufs['CO_NCM'].apply(get_sh6)
            df_imp_princ_ufs['SH6'] = df_imp_princ_ufs['CO_NCM'].apply(get_sh6)
            df_imp_comp_ufs['SH6'] = df_imp_comp_ufs['CO_NCM'].apply(get_sh6)

            # --- Filtra DFs de UF por mês ---
            df_exp_princ_ufs = df_exp_princ_ufs[df_exp_princ_ufs['CO_MES'].isin(meses_para_filtrar)]
            df_exp_comp_ufs = df_exp_comp_ufs[df_exp_comp_ufs['CO_MES'].isin(meses_para_filtrar)]
            df_imp_princ_ufs = df_imp_princ_ufs[df_imp_princ_ufs['CO_MES'].isin(meses_para_filtrar)]
            df_imp_comp_ufs = df_imp_comp_ufs[df_imp_comp_ufs['CO_MES'].isin(meses_para_filtrar)]

            # --- Filtra DFs de MG (para a UI) ---
            df_exp_princ_mg = df_exp_princ_ufs[df_exp_princ_ufs['SG_UF_NCM'] == 'MG']
            df_exp_comp_mg = df_exp_comp_ufs[df_exp_comp_ufs['SG_UF_NCM'] == 'MG']
            df_imp_princ_mg = df_imp_princ_ufs[df_imp_princ_ufs['SG_UF_NCM'] == 'MG']
            df_imp_comp_mg = df_imp_comp_ufs[df_imp_comp_ufs['SG_UF_NCM'] == 'MG']
            
            
            # --- Lógica de Loop (Agrupado vs Separado) ---
            if agrupado:
                nome_grupo = nome_agrupamento if (nome_agrupamento and nome_agrupamento.strip() != "") else ", ".join([p.split(' - ')[1] for p in produtos_para_agrupar_nomes])
                produtos_para_processar = [{
                    "nome": nome_grupo,
                    "codigos_sh2": codigos_sh2_selecionados,
                    "codigos_sh4": codigos_sh4_selecionados,
                    "codigos_sh6": codigos_sh6_selecionados,
                    "nomes_originais": produtos_para_agrupar_nomes 
                }]
            else:
                produtos_para_processar = []
                for nome_completo in sh2_selecionados_nomes:
                    produtos_para_processar.append({ "nome": nome_completo, "codigos_sh2": [nome_completo.split(" - ")[0]], "codigos_sh4": [], "codigos_sh6": [], "nomes_originais": [nome_completo] })
                for nome_completo in sh4_selecionados_nomes:
                    produtos_para_processar.append({ "nome": nome_completo, "codigos_sh2": [], "codigos_sh4": [nome_completo.split(" - ")[0]], "codigos_sh6": [], "nomes_originais": [nome_completo] })
                for nome_completo in sh6_selecionados_nomes:
                    produtos_para_processar.append({ "nome": nome_completo, "codigos_sh2": [], "codigos_sh4": [], "codigos_sh6": [nome_completo.split(" - ")[0]], "nomes_originais": [nome_completo] })
            
            # Loop principal de processamento
            for produto_info in produtos_para_processar:
                
                app = DocumentoApp(logo_path=logo_path_to_use)
                
                if agrupado:
                    st.subheader(f"Análise Agrupada de: {produto_info['nome']}")
                    nome_limpo_arquivo = sanitize_filename(produto_info['nome'])
                    titulo_doc = f"Briefing - {nome_limpo_arquivo} - {ano_principal}"
                    produto_nome_doc = f"de {produto_info['nome']}"
                else:
                    st.subheader(f"Análise de: {produto_info['nome']}")
                    nome_limpo_arquivo = sanitize_filename(produto_info['nome'].split(" - ")[1])
                    titulo_doc = f"Briefing - {nome_limpo_arquivo} - {ano_principal}"
                    produto_nome_doc = f"de {produto_info['nome'].split(' - ')[1]}"
                
                app.set_titulo(titulo_doc)

                # --- Lógica de Filtragem (para UI) ---
                filtro_sh2_mg = df_exp_princ_mg['SH2'].isin(produto_info['codigos_sh2'])
                filtro_sh4_mg = df_exp_princ_mg['SH4'].isin(produto_info['codigos_sh4'])
                filtro_sh6_mg = df_exp_princ_mg['SH6'].isin(produto_info['codigos_sh6'])
                df_exp_princ_f = df_exp_princ_mg[filtro_sh2_mg | filtro_sh4_mg | filtro_sh6_mg]

                filtro_sh2_comp_mg = df_exp_comp_mg['SH2'].isin(produto_info['codigos_sh2'])
                filtro_sh4_comp_mg = df_exp_comp_mg['SH4'].isin(produto_info['codigos_sh4'])
                filtro_sh6_comp_mg = df_exp_comp_mg['SH6'].isin(produto_info['codigos_sh6'])
                df_exp_comp_f = df_exp_comp_mg[filtro_sh2_comp_mg | filtro_sh4_comp_mg | filtro_sh6_comp_mg]
                
                if codigos_paises_selecionados:
                    df_exp_princ_f = df_exp_princ_f[df_exp_princ_f['CO_PAIS'].isin(codigos_paises_selecionados)]
                    df_exp_comp_f = df_exp_comp_f[df_exp_comp_f['CO_PAIS'].isin(codigos_paises_selecionados)]
                
                # --- UI: Tabela Exportação ---
                st.header("Principais Destinos (Exportação de MG)")
                exp_paises_princ = df_exp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False).reset_index()
                exp_paises_comp = df_exp_comp_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                # ... (resto da lógica da tabela da UI) ...
                exp_paises_princ['País'] = exp_paises_princ['CO_PAIS'].map(mapa_nomes_paises).fillna("Desconhecido")
                exp_paises_princ[f'Valor {ano_principal} (US$)'] = exp_paises_princ['VL_FOB']
                exp_paises_comp['País'] = exp_paises_comp['CO_PAIS'].map(mapa_nomes_paises).fillna("Desconhecido")
                exp_paises_comp[f'Valor {ano_comparacao} (US$)'] = exp_paises_comp['VL_FOB']
                exp_final = pd.merge(exp_paises_princ[['País', f'Valor {ano_principal} (US$)']], exp_paises_comp[['País', f'Valor {ano_comparacao} (US$)']], on="País", how="outer").fillna(0)
                exp_final['Variação %'] = 100 * (exp_final[f'Valor {ano_principal} (US$)'] - exp_final[f'Valor {ano_comparacao} (US$)']) / exp_final[f'Valor {ano_comparacao} (US$)']
                exp_final['Variação %'] = exp_final['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                exp_final[f'Valor {ano_principal}'] = exp_final[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                exp_final[f'Valor {ano_comparacao}'] = exp_final[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                df_display_exp = exp_final.sort_values(by=f'Valor {ano_principal} (US$)', ascending=False).reset_index(drop=True)
                st.dataframe(
                    df_display_exp[['País', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_paises),
                    hide_index=True,
                    use_container_width=True
                )
                
                # --- UI: Expander Exportação ---
                if agrupado and total_selecionado > 1:
                    with st.expander("Ver detalhamento de produtos por país (Exportação)"):
                        # ... (lógica do expander mantida) ...
                        top_paises_lista = df_display_exp['País'].head(top_n_paises).tolist()
                        def map_produto_selecionado(row):
                            if row['SH6'] in produto_info['codigos_sh6']: return mapa_sh6_nomes.get(row['SH6'], row['SH6'])
                            if row['SH4'] in produto_info['codigos_sh4']: return mapa_sh4_nomes.get(row['SH4'], row['SH4'])
                            if row['SH2'] in produto_info['codigos_sh2']: return mapa_sh2_nomes.get(row['SH2'], row['SH2'])
                            return None 
                        df_exp_princ_f['Produto'] = df_exp_princ_f.apply(map_produto_selecionado, axis=1)
                        df_exp_comp_f['Produto'] = df_exp_comp_f.apply(map_produto_selecionado, axis=1)
                        df_exp_princ_f_detalhe = df_exp_princ_f.dropna(subset=['Produto'])
                        df_exp_comp_f_detalhe = df_exp_comp_f.dropna(subset=['Produto'])
                        detalhe_exp_princ = df_exp_princ_f_detalhe.groupby(['CO_PAIS', 'Produto'])['VL_FOB'].sum().reset_index()
                        detalhe_exp_comp = df_exp_comp_f_detalhe.groupby(['CO_PAIS', 'Produto'])['VL_FOB'].sum().reset_index()
                        detalhe_exp_princ['País'] = detalhe_exp_princ['CO_PAIS'].map(mapa_nomes_paises)
                        detalhe_exp_comp['País'] = detalhe_exp_comp['CO_PAIS'].map(mapa_nomes_paises)
                        detalhe_exp_princ = detalhe_exp_princ.rename(columns={'VL_FOB': f'Valor {ano_principal} (US$)'})
                        detalhe_exp_comp = detalhe_exp_comp.rename(columns={'VL_FOB': f'Valor {ano_comparacao} (US$)'})
                        detalhe_exp_final = pd.merge(detalhe_exp_princ[['País', 'Produto', f'Valor {ano_principal} (US$)']], detalhe_exp_comp[['País', 'Produto', f'Valor {ano_comparacao} (US$)']], on=['País', 'Produto'], how='outer').fillna(0)
                        detalhe_exp_final = detalhe_exp_final[detalhe_exp_final['País'].isin(top_paises_lista)]
                        detalhe_exp_final['Variação %'] = 100 * (detalhe_exp_final[f'Valor {ano_principal} (US$)'] - detalhe_exp_final[f'Valor {ano_comparacao} (US$)']) / detalhe_exp_final[f'Valor {ano_comparacao} (US$)']
                        detalhe_exp_final['Variação %'] = detalhe_exp_final['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                        detalhe_exp_final[f'Valor {ano_principal}'] = detalhe_exp_final[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                        detalhe_exp_final[f'Valor {ano_comparacao}'] = detalhe_exp_final[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                        detalhe_exp_final = detalhe_exp_final.sort_values(by=['País', f'Valor {ano_principal} (US$)'], ascending=[True, False])
                        st.dataframe(detalhe_exp_final[['País', 'Produto', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']], hide_index=True, use_container_width=True)
                
                # --- UI: Tabela Importação ---
                st.header("Principais Origens (Importação de MG)")
                filtro_sh2_mg_imp = df_imp_princ_mg['SH2'].isin(produto_info['codigos_sh2'])
                filtro_sh4_mg_imp = df_imp_princ_mg['SH4'].isin(produto_info['codigos_sh4'])
                filtro_sh6_mg_imp = df_imp_princ_mg['SH6'].isin(produto_info['codigos_sh6'])
                df_imp_princ_f = df_imp_princ_mg[filtro_sh2_mg_imp | filtro_sh4_mg_imp | filtro_sh6_mg_imp]

                filtro_sh2_comp_mg_imp = df_imp_comp_mg['SH2'].isin(produto_info['codigos_sh2'])
                filtro_sh4_comp_mg_imp = df_imp_comp_mg['SH4'].isin(produto_info['codigos_sh4'])
                filtro_sh6_comp_mg_imp = df_imp_comp_mg['SH6'].isin(produto_info['codigos_sh6'])
                df_imp_comp_f = df_imp_comp_mg[filtro_sh2_comp_mg_imp | filtro_sh4_comp_mg_imp | filtro_sh6_comp_mg_imp]

                if codigos_paises_selecionados:
                    df_imp_princ_f = df_imp_princ_f[df_imp_princ_f['CO_PAIS'].isin(codigos_paises_selecionados)]
                    df_imp_comp_f = df_imp_comp_f[df_imp_comp_f['CO_PAIS'].isin(codigos_paises_selecionados)]

                imp_paises_princ = df_imp_princ_f.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False).reset_index()
                imp_paises_comp = df_imp_comp_f.groupby('CO_PAIS')['VL_FOB'].sum().reset_index()
                # ... (resto da lógica da tabela da UI) ...
                imp_paises_princ['País'] = imp_paises_princ['CO_PAIS'].map(mapa_nomes_paises).fillna("Desconhecido")
                imp_paises_princ[f'Valor {ano_principal} (US$)'] = imp_paises_princ['VL_FOB']
                imp_paises_comp['País'] = imp_paises_comp['CO_PAIS'].map(mapa_nomes_paises).fillna("Desconhecido")
                imp_paises_comp[f'Valor {ano_comparacao} (US$)'] = imp_paises_comp['VL_FOB']
                imp_final = pd.merge(imp_paises_princ[['País', f'Valor {ano_principal} (US$)']], imp_paises_comp[['País', f'Valor {ano_comparacao} (US$)']], on="País", how="outer").fillna(0)
                imp_final['Variação %'] = 100 * (imp_final[f'Valor {ano_principal} (US$)'] - imp_final[f'Valor {ano_comparacao} (US$)']) / imp_final[f'Valor {ano_comparacao} (US$)']
                imp_final['Variação %'] = imp_final['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                imp_final[f'Valor {ano_principal}'] = imp_final[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                imp_final[f'Valor {ano_comparacao}'] = imp_final[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                df_display_imp = imp_final.sort_values(by=f'Valor {ano_principal} (US$)', ascending=False).reset_index(drop=True)
                st.dataframe(
                    df_display_imp[['País', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']].head(top_n_paises),
                    hide_index=True,
                    use_container_width=True
                )
                
                # --- UI: Expander Importação ---
                if agrupado and total_selecionado > 1:
                    with st.expander("Ver detalhamento de produtos por país (Importação)"):
                        # ... (lógica do expander mantida) ...
                        top_paises_lista_imp = df_display_imp['País'].head(top_n_paises).tolist()
                        def map_produto_selecionado_imp(row):
                            if row['SH6'] in produto_info['codigos_sh6']: return mapa_sh6_nomes.get(row['SH6'], row['SH6'])
                            if row['SH4'] in produto_info['codigos_sh4']: return mapa_sh4_nomes.get(row['SH4'], row['SH4'])
                            if row['SH2'] in produto_info['codigos_sh2']: return mapa_sh2_nomes.get(row['SH2'], row['SH2'])
                            return None
                        df_imp_princ_f['Produto'] = df_imp_princ_f.apply(map_produto_selecionado_imp, axis=1)
                        df_imp_comp_f['Produto'] = df_imp_comp_f.apply(map_produto_selecionado_imp, axis=1)
                        df_imp_princ_f_detalhe = df_imp_princ_f.dropna(subset=['Produto'])
                        df_imp_comp_f_detalhe = df_imp_comp_f.dropna(subset=['Produto'])
                        detalhe_imp_princ = df_imp_princ_f_detalhe.groupby(['CO_PAIS', 'Produto'])['VL_FOB'].sum().reset_index()
                        detalhe_imp_comp = df_imp_comp_f_detalhe.groupby(['CO_PAIS', 'Produto'])['VL_FOB'].sum().reset_index()
                        detalhe_imp_princ['País'] = detalhe_imp_princ['CO_PAIS'].map(mapa_nomes_paises)
                        detalhe_imp_comp['País'] = detalhe_imp_comp['CO_PAIS'].map(mapa_nomes_paises)
                        detalhe_imp_princ = detalhe_imp_princ.rename(columns={'VL_FOB': f'Valor {ano_principal} (US$)'})
                        detalhe_imp_comp = detalhe_imp_comp.rename(columns={'VL_FOB': f'Valor {ano_comparacao} (US$)'})
                        detalhe_imp_final = pd.merge(detalhe_imp_princ[['País', 'Produto', f'Valor {ano_principal} (US$)']], detalhe_imp_comp[['País', 'Produto', f'Valor {ano_comparacao} (US$)']], on=['País', 'Produto'], how='outer').fillna(0)
                        detalhe_imp_final = detalhe_imp_final[detalhe_imp_final['País'].isin(top_paises_lista_imp)]
                        detalhe_imp_final['Variação %'] = 100 * (detalhe_imp_final[f'Valor {ano_principal} (US$)'] - detalhe_imp_final[f'Valor {ano_comparacao} (US$)']) / detalhe_imp_final[f'Valor {ano_comparacao} (US$)']
                        detalhe_imp_final['Variação %'] = detalhe_imp_final['Variação %'].replace([float('inf'), float('-inf')], 0).fillna(0).round(2)
                        detalhe_imp_final[f'Valor {ano_principal}'] = detalhe_imp_final[f'Valor {ano_principal} (US$)'].apply(formatar_valor)
                        detalhe_imp_final[f'Valor {ano_comparacao}'] = detalhe_imp_final[f'Valor {ano_comparacao} (US$)'].apply(formatar_valor)
                        detalhe_imp_final = detalhe_imp_final.sort_values(by=['País', f'Valor {ano_principal} (US$)'], ascending=[True, False])
                        st.dataframe(detalhe_imp_final[['País', 'Produto', f'Valor {ano_principal}', f'Valor {ano_comparacao}', 'Variação %']], hide_index=True, use_container_width=True)

                
                # --- NOVO: GERAÇÃO DE TEXTO PARA O DOCX ---
                
                # --- Filtros de Produto (para todos os DFs de UF) ---
                filtro_exp_sh2 = df_exp_princ_ufs['SH2'].isin(produto_info['codigos_sh2'])
                filtro_exp_sh4 = df_exp_princ_ufs['SH4'].isin(produto_info['codigos_sh4'])
                filtro_exp_sh6 = df_exp_princ_ufs['SH6'].isin(produto_info['codigos_sh6'])
                df_exp_princ_ufs_filtrado = df_exp_princ_ufs[filtro_exp_sh2 | filtro_exp_sh4 | filtro_exp_sh6]

                filtro_exp_comp_sh2 = df_exp_comp_ufs['SH2'].isin(produto_info['codigos_sh2'])
                filtro_exp_comp_sh4 = df_exp_comp_ufs['SH4'].isin(produto_info['codigos_sh4'])
                filtro_exp_comp_sh6 = df_exp_comp_ufs['SH6'].isin(produto_info['codigos_sh6'])
                df_exp_comp_ufs_filtrado = df_exp_comp_ufs[filtro_exp_comp_sh2 | filtro_exp_comp_sh4 | filtro_exp_comp_sh6]
                
                filtro_imp_sh2 = df_imp_princ_ufs['SH2'].isin(produto_info['codigos_sh2'])
                filtro_imp_sh4 = df_imp_princ_ufs['SH4'].isin(produto_info['codigos_sh4'])
                filtro_imp_sh6 = df_imp_princ_ufs['SH6'].isin(produto_info['codigos_sh6'])
                df_imp_princ_ufs_filtrado = df_imp_princ_ufs[filtro_imp_sh2 | filtro_imp_sh4 | filtro_imp_sh6]
                
                filtro_imp_comp_sh2 = df_imp_comp_ufs['SH2'].isin(produto_info['codigos_sh2'])
                filtro_imp_comp_sh4 = df_imp_comp_ufs['SH4'].isin(produto_info['codigos_sh4'])
                filtro_imp_comp_sh6 = df_imp_comp_ufs['SH6'].isin(produto_info['codigos_sh6'])
                df_imp_comp_ufs_filtrado = df_imp_comp_ufs[filtro_imp_comp_sh2 | filtro_imp_comp_sh4 | filtro_imp_comp_sh6]

                # --- Inicia Seção 1: Exportações ---
                app.nova_secao()
                app.adicionar_titulo("1. Exportações de Minas Gerais")

                # Parágrafo 1: Ranking Nacional
                ranking_exp_uf = df_exp_princ_ufs_filtrado.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(ascending=False)
                valor_total_br_exp = ranking_exp_uf.sum()
                valor_mg_exp = ranking_exp_uf.get('MG', 0)
                posicao_mg_exp = 0
                if valor_mg_exp > 0:
                    try:
                        posicao_mg_exp = ranking_exp_uf.index.get_loc('MG') + 1
                    except KeyError:
                        posicao_mg_exp = 0 # MG não exportou esse produto
                
                participacao_mg_exp = 0
                if valor_total_br_exp > 0:
                    participacao_mg_exp = (valor_mg_exp / valor_total_br_exp) * 100
                
                texto_exp_1 = f"Em {nome_periodo}, Minas Gerais foi o {posicao_mg_exp}º estado brasileiro exportador {produto_nome_doc}, com uma participação de {participacao_mg_exp:.2f}% nas exportações nacionais."
                app.adicionar_conteudo_formatado(texto_exp_1)

                # Parágrafo 2: Variação MG
                valor_mg_exp_comp = df_exp_comp_ufs_filtrado[df_exp_comp_ufs_filtrado['SG_UF_NCM'] == 'MG']['VL_FOB'].sum()
                dif_exp, tipo_dif_exp = calcular_diferenca_percentual(valor_mg_exp, valor_mg_exp_comp)
                texto_exp_2 = f"O estado exportou um montante de {formatar_valor(valor_mg_exp)}, apresentando {tipo_dif_exp} de {dif_exp:.1f}% em relação a {nome_periodo_comp}."
                app.adicionar_conteudo_formatado(texto_exp_2)

                # Parágrafo 3 e 4: Ranking Países
                df_exp_mg_filtrado = df_exp_princ_ufs_filtrado[df_exp_princ_ufs_filtrado['SG_UF_NCM'] == 'MG']
                ranking_paises_exp = df_exp_mg_filtrado.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
                total_mercados_exp = len(ranking_paises_exp)
                total_exp_mg = valor_mg_exp # Já calculado
                
                if total_mercados_exp > 0:
                    top_5_paises_exp = ranking_paises_exp.head(5)
                    lista_paises_texto = []
                    soma_top_5_exp = 0
                    for co_pais, valor in top_5_paises_exp.items():
                        nome_pais = mapa_nomes_paises.get(co_pais, "Desconhecido")
                        part = (valor / total_exp_mg) * 100
                        lista_paises_texto.append(f"{nome_pais} ({part:.2f}%)")
                        soma_top_5_exp += valor
                    
                    part_top_5_exp = (soma_top_5_exp / total_exp_mg) * 100
                    texto_exp_3 = f"As exportações {produto_nome_doc} de Minas Gerais atingiram {total_mercados_exp} mercados em {ano_principal}. Dentre esses, os maiores foram: {'; '.join(lista_paises_texto)}."
                    app.adicionar_conteudo_formatado(texto_exp_3)
                    texto_exp_4 = f"Juntos, esses cinco países foram responsáveis por {part_top_5_exp:.2f}% das exportações {produto_nome_doc} do estado."
                    app.adicionar_conteudo_formatado(texto_exp_4)

                # Parágrafo 5: Drill-Down de Produto
                nivel_detalhe = None
                mapa_detalhe = None
                if produto_info['codigos_sh6']: # Se selecionou SH6, não há drill-down
                    pass
                elif produto_info['codigos_sh4']: # Se selecionou SH4, detalha SH6
                    nivel_detalhe = 'SH6'
                    mapa_detalhe = mapa_sh6_nomes
                elif produto_info['codigos_sh2']: # Se selecionou SH2, detalha SH4
                    nivel_detalhe = 'SH4'
                    mapa_detalhe = mapa_sh4_nomes
                
                if nivel_detalhe and total_exp_mg > 0:
                    ranking_detalhe = df_exp_mg_filtrado.groupby(nivel_detalhe)['VL_FOB'].sum().sort_values(ascending=False).head(5)
                    lista_produtos_texto = []
                    for cod, valor in ranking_detalhe.items():
                        nome_prod = mapa_detalhe.get(cod, "Desconhecido")
                        part = (valor / total_exp_mg) * 100
                        lista_produtos_texto.append(f"{nome_prod} ({part:.2f}%)")
                    
                    texto_exp_5 = f"Em {ano_principal}, os principais produtos ({nivel_detalhe}) do setor {produto_nome_doc} exportados de Minas Gerais foram: {'; '.join(lista_produtos_texto)}."
                    app.adicionar_conteudo_formatado(texto_exp_5)

                # Parágrafo 6: Ranking Municípios (Impossível com estes dados)
                # (Omitido)


                # --- Inicia Seção 2: Importações ---
                app.nova_secao()
                app.adicionar_titulo("2. Importações de Minas Gerais")

                # Parágrafo 1: Ranking Nacional
                ranking_imp_uf = df_imp_princ_ufs_filtrado.groupby('SG_UF_NCM')['VL_FOB'].sum().sort_values(ascending=False)
                valor_total_br_imp = ranking_imp_uf.sum()
                valor_mg_imp = ranking_imp_uf.get('MG', 0)
                posicao_mg_imp = 0
                if valor_mg_imp > 0:
                    try:
                        posicao_mg_imp = ranking_imp_uf.index.get_loc('MG') + 1
                    except KeyError:
                        posicao_mg_imp = 0
                
                participacao_mg_imp = 0
                if valor_total_br_imp > 0:
                    participacao_mg_imp = (valor_mg_imp / valor_total_br_imp) * 100
                
                texto_imp_1 = f"Em {nome_periodo}, Minas Gerais foi o {posicao_mg_imp}º estado brasileiro importador {produto_nome_doc}, com uma participação de {participacao_mg_imp:.2f}% nas importações nacionais."
                app.adicionar_conteudo_formatado(texto_imp_1)

                # Parágrafo 2: Variação MG
                valor_mg_imp_comp = df_imp_comp_ufs_filtrado[df_imp_comp_ufs_filtrado['SG_UF_NCM'] == 'MG']['VL_FOB'].sum()
                dif_imp, tipo_dif_imp = calcular_diferenca_percentual(valor_mg_imp, valor_mg_imp_comp)
                texto_imp_2 = f"O estado importou um montante de {formatar_valor(valor_mg_imp)}, apresentando {tipo_dif_imp} de {dif_imp:.1f}% em relação a {nome_periodo_comp}."
                app.adicionar_conteudo_formatado(texto_imp_2)

                # Parágrafo 3 e 4: Ranking Países
                df_imp_mg_filtrado = df_imp_princ_ufs_filtrado[df_imp_princ_ufs_filtrado['SG_UF_NCM'] == 'MG']
                ranking_paises_imp = df_imp_mg_filtrado.groupby('CO_PAIS')['VL_FOB'].sum().sort_values(ascending=False)
                total_mercados_imp = len(ranking_paises_imp)
                total_imp_mg = valor_mg_imp
                
                if total_mercados_imp > 0:
                    top_5_paises_imp = ranking_paises_imp.head(5)
                    lista_paises_texto_imp = []
                    soma_top_5_imp = 0
                    for co_pais, valor in top_5_paises_imp.items():
                        nome_pais = mapa_nomes_paises.get(co_pais, "Desconhecido")
                        part = (valor / total_imp_mg) * 100
                        lista_paises_texto_imp.append(f"{nome_pais} ({part:.2f}%)")
                        soma_top_5_imp += valor
                    
                    part_top_5_imp = (soma_top_5_imp / total_imp_mg) * 100
                    texto_imp_3 = f"As importações mineiras {produto_nome_doc} tiveram origem em {total_mercados_imp} mercados em {ano_principal}. Dentre esses, os maiores foram: {'; '.join(lista_paises_texto_imp)}."
                    app.adicionar_conteudo_formatado(texto_imp_3)
                    texto_imp_4 = f"Juntos, esses cinco países foram responsáveis por {part_top_5_imp:.2f}% das importações {produto_nome_doc} do estado."
                    app.adicionar_conteudo_formatado(texto_imp_4)

                # Parágrafo 5: Drill-Down de Produto
                nivel_detalhe_imp = None
                mapa_detalhe_imp = None
                if produto_info['codigos_sh6']: 
                    pass
                elif produto_info['codigos_sh4']: 
                    nivel_detalhe_imp = 'SH6'
                    mapa_detalhe_imp = mapa_sh6_nomes
                elif produto_info['codigos_sh2']:
                    nivel_detalhe_imp = 'SH4'
                    mapa_detalhe_imp = mapa_sh4_nomes
                
                if nivel_detalhe_imp and total_imp_mg > 0:
                    ranking_detalhe_imp = df_imp_mg_filtrado.groupby(nivel_detalhe_imp)['VL_FOB'].sum().sort_values(ascending=False).head(5)
                    lista_produtos_texto_imp = []
                    for cod, valor in ranking_detalhe_imp.items():
                        nome_prod = mapa_detalhe_imp.get(cod, "Desconhecido")
                        part = (valor / total_imp_mg) * 100
                        lista_produtos_texto_imp.append(f"{nome_prod} ({part:.2f}%)")
                    
                    texto_imp_5 = f"Em {ano_principal}, os principais produtos ({nivel_detalhe_imp}) do setor {produto_nome_doc} importados por Minas Gerais foram: {'; '.join(lista_produtos_texto_imp)}."
                    app.adicionar_conteudo_formatado(texto_imp_5)

                # Parágrafo 6: Ranking Municípios (Impossível com estes dados)
                # (Omitido)

                # --- FIM DA GERAÇÃO DE TEXTO ---

                # Salva o documento no state
                file_bytes, file_name = app.finalizar_documento()
                st.session_state.arquivos_gerados_produto.append({"name": file_name, "data": file_bytes})
            
            # Limpa DFs grandes da memória
            del df_exp_princ_ufs, df_exp_comp_ufs, df_imp_princ_ufs, df_imp_comp_ufs
            del df_exp_princ_mg, df_exp_comp_mg, df_imp_princ_mg, df_imp_comp_mg

        except Exception as e:
            st.error(f"Ocorreu um erro inesperado durante a análise de produto:")
            st.exception(e)

# --- Bloco de exibição de Download (COM LÓGICA DE ZIP) ---
if st.session_state.arquivos_gerados_produto:
    st.header("4. Relatórios Gerados")
    st.info("Clique para baixar os relatórios. Eles permanecerão aqui até que você gere um novo relatório.")
    
    if len(st.session_state.arquivos_gerados_produto) > 1:
        st.subheader("Pacote de Relatórios (ZIP)")
        
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for arquivo in st.session_state.arquivos_gerados_produto:
                zip_file.writestr(arquivo["name"], arquivo["data"])
        
        zip_bytes = zip_buffer.getvalue()
        
        st.download_button(
            label=f"Baixar todos os {len(st.session_state.arquivos_gerados_produto)} relatórios (.zip)",
            data=zip_bytes,
            file_name=f"Briefings_Produtos_{ano_principal}.zip",
            mime="application/zip",
            key="download_zip_produto"
        )
        
    elif len(st.session_state.arquivos_gerados_produto) == 1:
        st.subheader("Relatório Gerado")
        arquivo = st.session_state.arquivos_gerados_produto[0] 
        st.download_button(
            label=f"Baixar Relatório ({arquivo['name']})",
            data=arquivo["data"], 
            file_name=arquivo["name"],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{arquivo['name']}"
        )

# --- Bloco de Rodapé ---
st.divider() 

col1, col2 = st.columns([0.3, 0.7], vertical_alignment="center") 

with col1:
    logo_footer_path = "AEST Sede.png"
    if os.path.exists(logo_footer_path):
        st.image(logo_footer_path, width=150)
    else:
        st.caption("Logo AEST não encontrada.")

with col2:
    st.caption("Desenvolvido por Aest - Dados e Subsecretaria de Promoção de Investimentos e Cadeias Produtivas")
